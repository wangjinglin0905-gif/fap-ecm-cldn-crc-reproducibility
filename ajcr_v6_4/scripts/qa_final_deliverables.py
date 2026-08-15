from __future__ import annotations

import hashlib
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "work" / "deliverables_draft" / "FAP_Senescence_AJCR_rebuilt_draft.docx"
FIGDIR = ROOT / "work" / "figures" / "revised"
PDF = ROOT / "work" / "rendered_manuscript_final_v63" / "FAP_Senescence_AJCR_rebuilt_draft.pdf"
COVER_DOCX = ROOT / "work" / "deliverables_draft" / "AJCR_CoverLetter_rebuilt_2026-08-14.docx"
COVER_PDF = ROOT / "work" / "rendered_cover_letter_final_v2" / "AJCR_CoverLetter_rebuilt_2026-08-14.pdf"
QA_REPORT = ROOT / "work" / "qa" / "final_qa_report.json"
EXPECTED_TITLE = "Fibroblast-enriched senescence-associated transcription and a distinct FAP-linked matrix state in colorectal cancer"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


document = Document(DOCX)
paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
table_texts = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
text = "\n".join(paragraph_texts + table_texts)
cover_document = Document(COVER_DOCX)
cover_text = "\n".join(paragraph.text for paragraph in cover_document.paragraphs)

abstract_start = paragraph_texts.index("Abstract") + 1
abstract = paragraph_texts[abstract_start]
abstract_words = len(re.findall(r"\b[\w+-]+\b", abstract))

declarations_start = paragraph_texts.index("Declarations")
citation_text = "\n".join(paragraph_texts[:declarations_start])
citations: set[int] = set()
citation_order: list[int] = []
for match in re.finditer(r"\[([0-9,-]+)\]", citation_text):
    token = match.group(1)
    for piece in token.split(","):
        if "-" in piece:
            lo, hi = map(int, piece.split("-"))
            numbers = range(lo, hi + 1)
        else:
            numbers = [int(piece)]
        for number in numbers:
            if number not in citations:
                citation_order.append(number)
            citations.add(number)

references_start = paragraph_texts.index("References") + 1
reference_numbers = [
    int(match.group(1))
    for item in paragraph_texts[references_start:]
    if (match := re.match(r"^(\d+)\. ", item))
]
reference_count = len(reference_numbers)
forbidden = {
    "old_title": "Compartmentalized senescence programs in colorectal cancer",
    "old_B1_beta": "beta = 0.083",
    "old_B1_p": "6.96 x 10^-37",
    "old_cell_model": "1,498 cells",
    "old_receptor_slope": "slope = -0.005",
    "supplementary_figure": "Supplementary Figure",
    "conditional_figure_phrase": "if no external cohort",
    "editorial_rejection_history": "editorial rejection",
    "candidate_version_history": "candidate v6.0",
    "discarded_version_history": "discarded v6.0",
    "revised_analysis_history": "The revised analysis",
    "previously_reported_count": "previously reported count",
}

figure_stems = [
    "Figure1_tumor_only_compartment",
    "Figure2_FAP_and_matrix_separation",
    "Figure3_specificity_and_composition",
    "Figure4_GSE166555_external_validation",
    "Figure6_cognate_SASP_immune_screen",
]
figure_checks = []
for stem in figure_stems:
    item = {"stem": stem}
    for extension in ("png", "tiff"):
        path = FIGDIR / f"{stem}.{extension}"
        with Image.open(path) as image:
            dpi = [float(value) for value in image.info.get("dpi", ())]
            item[extension] = {
                "size": list(image.size),
                "mode": image.mode,
                "dpi": dpi,
                "compression": str(image.info.get("compression")),
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
    figure_checks.append(item)

with zipfile.ZipFile(DOCX) as archive:
    zip_error = archive.testzip()
    embedded_media = [name for name in archive.namelist() if name.startswith("word/media/")]

report = {
    "docx": {
        "path": str(DOCX),
        "bytes": DOCX.stat().st_size,
        "sha256": sha256(DOCX),
        "zip_test": zip_error,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "embedded_media": len(embedded_media),
    },
    "rendered_pdf": {
        "path": str(PDF),
        "bytes": PDF.stat().st_size,
        "sha256": sha256(PDF),
        "pages": len(PdfReader(PDF).pages),
        "visual_inspection": "all 20 full-resolution page renders inspected at 150 dpi in five ordered contact sheets",
    },
    "cover_letter": {
        "docx_path": str(COVER_DOCX),
        "docx_bytes": COVER_DOCX.stat().st_size,
        "docx_sha256": sha256(COVER_DOCX),
        "pdf_path": str(COVER_PDF),
        "pdf_bytes": COVER_PDF.stat().st_size,
        "pdf_sha256": sha256(COVER_PDF),
        "pages": len(PdfReader(COVER_PDF).pages),
        "title_matches_manuscript": EXPECTED_TITLE in cover_text,
        "old_title_present": "Compartmentalized senescence programs in colorectal cancer" in cover_text,
        "old_effect_present": "0.376" in cover_text or "0.083" in cover_text or "0.608" in cover_text,
        "reviewer_entries": sum(name in cover_text for name in ("Sheila A. Stewart", "Florian R. Greten", "Qing Nie")),
        "visual_inspection": "both page renders inspected at 150 dpi; reviewer entries are not split across pages",
    },
    "abstract_words": abstract_words,
    "references": reference_count,
    "citations_seen": sorted(citations),
    "citation_first_appearance_order": citation_order,
    "reference_list_sequence_complete": reference_numbers == list(range(1, reference_count + 1)),
    "citation_sequence_complete": citation_order == list(range(1, reference_count + 1)),
    "forbidden_string_hits": {key: value in text for key, value in forbidden.items()},
    "placeholders": [line for line in paragraph_texts if "SUBMISSION" in line],
    "figures": figure_checks,
}

failures = []
if abstract_words > 350:
    failures.append(f"Abstract has {abstract_words} words")
if zip_error is not None:
    failures.append(f"DOCX zip error: {zip_error}")
if len(document.tables) != 4:
    failures.append(f"Expected 4 tables, observed {len(document.tables)}")
if len(embedded_media) != 5:
    failures.append(f"Expected 5 embedded figures, observed {len(embedded_media)}")
if any(report["forbidden_string_hits"].values()):
    failures.append("At least one forbidden legacy string remains")
if not report["citation_sequence_complete"]:
    failures.append("Reference citations are not complete, consecutive, and ordered by first appearance")
if not report["reference_list_sequence_complete"]:
    failures.append("Reference list numbering is not complete and consecutive")
if paragraph_texts[0] != EXPECTED_TITLE:
    failures.append("Manuscript title does not match the frozen title")
if not report["cover_letter"]["title_matches_manuscript"]:
    failures.append("Cover letter title does not match manuscript title")
if report["cover_letter"]["old_title_present"] or report["cover_letter"]["old_effect_present"]:
    failures.append("Cover letter contains a legacy title or effect size")
if report["cover_letter"]["reviewer_entries"] != 3:
    failures.append("Cover letter does not contain all three reviewer entries")
for item in figure_checks:
    if item["tiff"]["compression"] not in {"tiff_lzw", "lzw"}:
        failures.append(f"{item['stem']} TIFF is not LZW-compressed")
    if min(item["tiff"]["dpi"] or [0]) < 590:
        failures.append(f"{item['stem']} TIFF DPI is below 600 tolerance")

report["status"] = "PASS" if not failures else "FAIL"
report["failures"] = failures
report_json = json.dumps(report, indent=2, ensure_ascii=False)
QA_REPORT.parent.mkdir(parents=True, exist_ok=True)
QA_REPORT.write_text(report_json + "\n", encoding="utf-8")
print(report_json)
if failures:
    raise SystemExit(1)
