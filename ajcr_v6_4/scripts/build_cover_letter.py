from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "work" / "deliverables_draft"
OUT.mkdir(parents=True, exist_ok=True)
FONT = "Arial"


def style_run(run, size=11, bold=False, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(document, text="", *, align=None, space_after=6, first_line=False):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.08
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    if text:
        style_run(paragraph.add_run(text))
    return paragraph


def build():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)

    title = add_paragraph(document, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    style_run(title.add_run("Cover Letter"), size=14, bold=True)

    add_paragraph(document, "14 August 2026", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(document, "Editor-in-Chief\nAmerican Journal of Cancer Research", space_after=10)
    add_paragraph(document, "Dear Editor-in-Chief,", space_after=8)

    p = add_paragraph(document, first_line=True)
    style_run(p.add_run("We are pleased to submit our original research article, \""))
    style_run(
        p.add_run(
            "Fibroblast-enriched senescence-associated transcription and a distinct "
            "FAP-linked matrix state in colorectal cancer"
        ),
        italic=True,
    )
    style_run(p.add_run("\", for consideration in the American Journal of Cancer Research."))

    add_paragraph(
        document,
        "The study addresses a question that bulk-tissue association alone cannot resolve: how senescence-associated "
        "transcription is distributed across colorectal cancer tumour compartments, whether it is concentrated within "
        "FAP-detected fibroblasts, and whether it follows the same biological axis as the FAP-associated extracellular-matrix "
        "program. We integrated two bulk cohorts, tumour-restricted single-cell analysis, an independent single-cell cohort "
        "and CPTAC colon proteomics, while separating overlapping score genes and treating the patient as the inferential unit.",
        first_line=True,
    )

    add_paragraph(
        document,
        "In GSE132465, the frozen SenMayo score was higher in tumour fibroblast-lineage than paired epithelial cells in all "
        "23 patients (median standardized difference 0.466, 95% bootstrap CI 0.410–0.553; exact P=2.38×10⁻⁷) and exceeded "
        "5,000 expression-matched null signatures. The direction recurred in 9 of 10 eligible patients in GSE166555 "
        "(P=0.0195). Depth- and subtype-adjusted analyses placed this signal at the fibroblast-compartment level rather than "
        "within the FAP-detected subgroup. In contrast, FAP retained a strong association with matrix expression in tumour "
        "fibroblasts and with an FAP-excluded ECM protein score. Bulk sensitivity analyses further showed that the very large "
        "marginal FAP–matrix correlations were substantially attenuated by MCP-counter and EPIC composition adjustment. "
        "Together, the results delineate two stromal axes and define a focused path for multi-marker tissue and functional validation.",
        first_line=True,
    )

    add_paragraph(document, "The main contributions are:", space_after=3)
    contributions = [
        "a tumour-only, patient-paired analysis with an independently evaluated compartment direction;",
        "a frozen, overlap-controlled SenMayo score and expression-matched null analysis;",
        "depth-, subtype- and composition-aware separation of senescence-associated transcription from the FAP–matrix state; and",
        "transparent reporting of receptor, immune-screen and prognostic boundaries, with reproducible figures and source outputs.",
    ]
    for item in contributions:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(2)
        style_run(p.add_run(item))

    add_paragraph(
        document,
        "We believe the manuscript will interest readers studying tumour stroma, cellular-state inference, colorectal cancer "
        "heterogeneity and the interpretation of public multi-omic data. The framing is intentionally evidence-matched: "
        "transcriptomic scores localize a reproducible state, while senescent-cell identity and functional communication are "
        "reserved for orthogonal validation.",
        first_line=True,
    )

    reviewer_heading = add_paragraph(document, "Suggested reviewers:", space_after=3)
    reviewer_heading.paragraph_format.keep_with_next = True
    reviewers = [
        (
            "Sheila A. Stewart, PhD",
            "Gerty Cori Professor and Vice Chair, Department of Cell Biology & Physiology; Professor of Medicine, "
            "Washington University School of Medicine, St Louis, MO, USA",
            "sheila.stewart@wustl.edu",
            "cellular senescence, stromal ageing and cancer-associated fibroblasts",
        ),
        (
            "Florian R. Greten, MD",
            "Director, Georg-Speyer-Haus Institute for Tumor Biology and Experimental Therapy, Frankfurt am Main, Germany",
            "greten@gsh.uni-frankfurt.de",
            "colorectal cancer microenvironment and stromal–immune interactions",
        ),
        (
            "Qing Nie, PhD",
            "Distinguished Professor of Mathematics, Developmental & Cell Biology and Systems Biology, "
            "University of California, Irvine, CA, USA",
            "qnie@uci.edu",
            "computational systems biology and single-cell cell–cell interaction inference",
        ),
    ]
    for index, (name, affiliation, email, expertise) in enumerate(reviewers, start=1):
        p = add_paragraph(document, space_after=3)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.keep_together = True
        style_run(p.add_run(f"{index}. {name}. "), bold=True)
        style_run(p.add_run(f"{affiliation}. {email}. Expertise: {expertise}."))

    add_paragraph(
        document,
        "All authors have read and approved the manuscript. The work has not been published and is not under consideration "
        "elsewhere. The authors declare no competing interests. Only publicly available, de-identified data were analysed, and "
        "no new human or animal specimens were collected.",
        first_line=True,
    )

    add_paragraph(document, "Thank you for your consideration.", space_after=10)
    add_paragraph(document, "Sincerely,", space_after=8)
    add_paragraph(document, "Jinglin Wang, MD\nDepartment of Geriatrics\nThe Affiliated Hospital of Guizhou Medical University\nGuiyang, China", space_after=8)
    add_paragraph(document, "Corresponding author: Juan Yang, MD, PhD\nEmail: yj63yj63@163.com")

    path = OUT / "AJCR_CoverLetter_rebuilt_2026-08-14.docx"
    document.save(path)
    print(f"Wrote {path}")
    return path


if __name__ == "__main__":
    build()
