from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "work"
OUT = WORK / "deliverables_draft"
OUT.mkdir(parents=True, exist_ok=True)

FONT = "Franklin Gothic Book"
FIGURE_DIR = WORK / "figures" / "revised"
EXT_STATS = WORK / "analysis" / "GSE166555_senescence" / "validation_summary.json"
EXT_FIGURE = FIGURE_DIR / "Figure4_GSE166555_external_validation.png"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    for style_name, size in (("Title", 16), ("Heading 1", 14), ("Heading 2", 12)):
        style = document.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0


def add_body(document: Document, text: str, *, align=None, keep=False):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.keep_together = keep
    paragraph.add_run(text)
    return paragraph


def add_inline_label(document: Document, label: str, text: str):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(text)
    return paragraph


def add_placeholder(document: Document, text: str):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.highlight_color = 7  # yellow
    return paragraph


def add_authors(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors = [
        ("Jinglin Wang", "1,2"), ("Pan Sun", "1"), ("Lei Xu", "1"),
        ("Yao Wang", "1"), ("Wei Zhu", "1"), ("Yan Zhang", "1"),
        ("Yao Tang", "1"), ("Juan Yang", "1,*"),
    ]
    for index, (name, affiliation) in enumerate(authors):
        if index:
            paragraph.add_run(", ")
        paragraph.add_run(name)
        superscript = paragraph.add_run(affiliation)
        superscript.font.superscript = True


def add_abstract(document: Document, external: dict | None) -> str:
    external_clause = ""
    if external:
        sen = external["SenMayo119"]
        external_clause = (
            f" The direction was externally evaluated in GSE166555: {sen['concordant']}/"
            f"{sen['n']} eligible patients had higher fibroblast scores "
            f"(median paired difference {sen['median_difference']:.3f}, exact P = {sen['p_exact']:.3g})."
        )
    abstract = (
        "Senescence-associated transcription is detectable in colorectal cancer (CRC), but its cellular distribution "
        "and relationship to fibroblast activation protein (FAP)-linked matrix remodelling remain uncertain. We integrated "
        "TCGA-COAD/READ and GSE39582 bulk transcriptomes, GSE132465 single-cell RNA sequencing, an external GSE166555 "
        "single-cell evaluation, and CPTAC colon proteomics. Analyses used a 119-gene SenMayo score after removal of FAP-score "
        "overlap, tumour-only cell sets, patient-paired inference, sequencing-depth and fibroblast-subtype adjustment, "
        "expression-matched null gene sets, and multiple fibroblast-abundance proxies. In GSE132465, SenMayo scores were higher "
        "in tumour fibroblast-lineage than epithelial cells in all 23 patients (median paired standardized difference 0.466, "
        "95% bootstrap CI 0.410-0.553; exact P = 2.38 x 10^-7), while MKI67 showed the opposite direction in 21 of 23 patients. "
        "The SenMayo contrast exceeded 5,000 expression-matched null signatures (empirical P = 2.00 x 10^-4)."
        f"{external_clause}"
        " Within tumour fibroblasts, the adjusted FAP-detection coefficient for SenMayo was 0.0238 "
        "(95% CI -0.0051 to 0.0527; P = 0.106); the SASP score and individual CDKN2A, CDKN2B, CDKN1A and LMNB1 "
        "detection models gave concordant null estimates. By contrast, patient-level fibroblast FAP tracked matrix4 "
        "(rho = 0.649, P = 8.03 x 10^-4), whereas FAP-SenMayo and SenMayo-matrix4 correlations were weaker. Bulk "
        "FAP13-matrix4 correlations attenuated markedly under MCP-counter/EPIC composition control, while CPTAC retained "
        "FAP-ECM protein covariation. These data delineate a reproducible fibroblast-enriched senescence-associated "
        "transcriptional pattern and a distinct, abundance-linked FAP-matrix state. The evidence localizes transcriptional "
        "features; cellular senescence and its functional consequences require orthogonal tissue and perturbation studies."
    )
    paragraph = document.add_paragraph()
    paragraph.add_run(abstract)
    return abstract


def add_table(document: Document, title: str, headers: list[str], rows: list[list[str]], note: str | None = None) -> None:
    document.add_page_break()
    caption = document.add_paragraph()
    caption.paragraph_format.keep_with_next = True
    run = caption.add_run(title)
    run.bold = True
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9EAF0")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            row.cells[index].text = str(value)
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = FONT
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
                    run.font.size = Pt(10)
    if note:
        paragraph = document.add_paragraph()
        paragraph.add_run("Note: ").bold = True
        paragraph.add_run(note)


def add_figure(document: Document, label: str, path: Path) -> None:
    document.add_page_break()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(label).bold = True
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(path), width=Inches(6.55))


def build() -> Path:
    external = json.loads(EXT_STATS.read_text(encoding="utf-8")) if EXT_STATS.exists() else None
    document = Document()
    configure_document(document)

    title = "Fibroblast-enriched senescence-associated transcription and a distinct FAP-linked matrix state in colorectal cancer"
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.keep_with_next = True
    title_paragraph.paragraph_format.space_after = Pt(3)
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = FONT
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    title_run.font.size = Pt(16)
    add_authors(document)
    add_body(document, "1 Department of Geriatrics, The Affiliated Hospital of Guizhou Medical University, Guiyang, China", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body(document, "2 Department of Digestive Disease, Guizhou Provincial People's Hospital, Guiyang, China", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body(document, "* Corresponding author: Juan Yang, MD, PhD; Email: yj63yj63@163.com; ORCID iD: 0009-0001-2081-6950", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body(document, "Running title: Stromal senescence-associated transcription in CRC", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_placeholder(document, "SUBMISSION CHECK: add institutional email addresses for all co-authors if required by the AJCR submission portal.")

    document.add_heading("Abstract", level=1)
    abstract = add_abstract(document, external)
    add_inline_label(
        document,
        "Keywords: ",
        "colorectal cancer; cellular senescence; fibroblast activation protein; cancer-associated fibroblast; "
        "senescence-associated secretory phenotype; extracellular matrix; single-cell RNA sequencing; composition adjustment"
    )

    document.add_heading("1. Introduction", level=1)
    add_body(document, "Colorectal cancer (CRC) remains a major cause of cancer morbidity and mortality worldwide [1]. Cancer-associated fibroblasts (CAFs) are heterogeneous components of the tumour microenvironment that shape extracellular matrix (ECM), paracrine and immune states [2]. Fibroblast activation protein (FAP) marks an activated stromal population in many malignancies, but FAP expression does not by itself define a single functional CAF state [3]. In CRC, stromal transcripts can dominate molecular classification and prognostic associations [4,5], while spatial analysis has linked FAP-positive fibroblasts with macrophage-rich niches [6]. These observations make it important to distinguish cell abundance, lineage-associated transcription and inferred communication.")
    add_body(document, "SenMayo was developed as a cross-tissue senescence-associated transcriptional panel [7]. Cellular senescence itself is a durable growth-arrest phenotype for which tissue assessment requires convergent markers and functional context [8,9]; loss of lamin B1 is one candidate feature rather than a universal criterion [10]. The senescence-associated secretory phenotype (SASP) is likewise context dependent and overlaps inflammatory and wound-response programs [11]. Transcriptomic scores can therefore localize senescence-associated features, whereas senescent-cell identity requires orthogonal evidence.")
    add_body(document, "Studies of senescent myofibroblasts in pancreatic cancer [12] and the core senescence phenotype of primary human colon fibroblasts [13] illustrate the value of convergent phenotyping. Two recent CRC studies extend this principle mechanistically. Ge et al. combined single-cell classification, tissue assessment and functional experiments to link senescent fibroblasts with CD8-positive T-cell dysfunction through lipid transfer and peroxidation [14]. Yang et al. integrated computational, spatial, organoid and genetic models to show that macrophage-induced senescent CAFs promoted SASP-mediated chemoresistance through an IL1B-IL1R1 axis [15]. The remaining question is how senescence-associated transcription is distributed across CRC tumour compartments, whether it is concentrated specifically in FAP-detected fibroblasts, and whether it follows the same biological axis as the FAP-associated matrix program.")
    add_body(document, "We addressed these questions by rebuilding the analysis around tumour-only, patient-paired single-cell comparisons and explicit negative controls. The primary hypotheses were that senescence-associated scores are unevenly distributed between fibroblast-lineage and epithelial compartments, and that any FAP-specific enrichment remains after adjustment for sequencing depth and fibroblast subtype. We then tested whether senescence-associated and FAP-matrix signals remained coupled under expression-matched and composition-aware analyses. Receptor and immune-communication analyses were retained only to define mechanistic boundaries.")

    document.add_heading("2. Materials and methods", level=1)
    document.add_heading("2.1 Study design and ethics", level=2)
    add_body(document, "This study was a secondary analysis of publicly available, de-identified human CRC data. No participant was recruited and no new human or animal specimen was collected. The analyses were exploratory and were not preregistered. Gene-set definitions and the GSE132465 primary analysis were frozen before analysis of the external GSE166555 cohort. Results are reported as molecular associations rather than causal or clinical evidence.")

    document.add_heading("2.2 Data sources and analysis populations", level=2)
    add_body(document, "Bulk RNA-sequencing data comprised 380 barcode-reviewed primary TCGA-COAD/READ tumours and 566 tumour samples from GSE39582 after exclusion of 19 non-tumour mucosa arrays [16]. GSE132465 included 63,689 annotated cells from 23 CRC patients [17]. Because the source object also contained normal tissue, the primary single-cell universe was restricted to source-annotated tumour cells: 1,501 fibroblast-lineage cells (Myofibroblasts and Stromal 1-3) and 17,469 epithelial cells. GSE166555 contained 68,702 annotated cells from 12 patients [18]; tumour CAF-like cells were defined by source annotations containing FBs, CAFs or MyoFBs, and patients required at least 20 CAF-like and 100 epithelial cells. CPTAC colon proteomics comprised 97 tumours with relevant proteins available through cBioPortal [19-21].")

    document.add_heading("2.3 Gene-set definitions and circularity control", level=2)
    add_body(document, "Scores were means of gene-wise standardized expression unless stated otherwise. FAP13 comprised FAP, POSTN, THY1, PDPN, TAGLN, ACTA2, MMP2, MMP9, CXCL12, TGFB1, INHBA, WNT2 and WNT5A. matrix4 comprised COL1A1, COL1A2, COL3A1 and FN1; receptor2 comprised SDC4 and CD44. The original SenMayo source list contained 125 genes [7]. CXCL12, MMP2, MMP9 and WNT2 overlapped FAP13 and were removed, leaving 121 candidates; BEX3 and CCL3L1 were absent from GSE132465, yielding a frozen 119-gene score. The SASP25 score comprised IL6, CXCL8, IL1A, IL1B, CCL2, CCL5, CXCL1, CXCL2, CXCL3, CXCL10, MMP1, MMP3, MMP9, MMP10, MMP13, SERPINE1, PLAU, TIMP2, VEGFA, GDF15, IGFBP3, TNF, CSF2, HGF and FAS. Score overlap and represented-gene counts were audited before inference.")

    document.add_heading("2.4 GSE132465 tumour-compartment analysis", level=2)
    add_body(document, "Raw UMI counts were normalized per cell as log1p(counts per 10,000). For each patient, mean SenMayo119, SASP25 and MKI67 values were summarized separately in tumour fibroblast-lineage and epithelial cells. Paired differences were tested with exact two-sided Wilcoxon signed-rank tests; 95% confidence intervals for median paired differences were obtained by 10,000 patient-level bootstrap resamples. A sensitivity analysis required at least 20 cells in both compartments. The patient, rather than each cell, was the inferential unit for paired contrasts, consistent with the hierarchical sampling structure of single-cell data [22,23].")

    document.add_heading("2.5 FAP-specific inference within tumour fibroblasts", level=2)
    add_body(document, "FAP-detected cells were defined as having at least one FAP UMI; this label is used instead of FAP-positive because transcript detection is not equivalent to protein positivity. The primary linear mixed model was score ~ FAP detection + MKI67 z score + log1p(nCount_RNA) + fibroblast subtype + (1|patient). Satterthwaite degrees of freedom were used for gene-set outcomes. Detection of CDKN2A, CDKN2B, CDKN1A, LMNB1 and MKI67 was analysed by logistic mixed models with log UMI depth, subtype and patient random intercept. Effects are reported with 95% confidence intervals. Unadjusted detection fractions were descriptive only.")

    document.add_heading("2.6 External single-cell evaluation in GSE166555", level=2)
    add_body(document, "Raw counts and metadata were downloaded from GEO. Counts were summed within each tumour patient-compartment group, converted to log1p counts per million, and standardized gene-wise across eligible patient-compartment pseudobulks. The frozen 119-gene SenMayo set was used without re-selection. Patient-paired fibroblast-epithelial differences were assessed with the same exact Wilcoxon and bootstrap procedures. This cohort was an external evaluation of the compartment direction, not a validation of cellular senescence.")

    document.add_heading("2.7 Expression-matched signature specificity", level=2)
    add_body(document, "To test whether lineage differences reflected generic expression characteristics, 5,000 random gene sets were matched to SenMayo119 or SASP25 on mean expression, standard deviation and detection rate in GSE132465. For each draw, the mean patient-level fibroblast-minus-epithelial difference was recomputed. The upper-tail empirical P value was (1 + number of null statistics at least as large as observed)/(5,000 + 1). This null addresses transcript abundance and variability, not all lineage biology.")

    document.add_heading("2.8 Bulk composition sensitivity", level=2)
    add_body(document, "Within TCGA, partial Spearman correlations were estimated after rank-transforming the molecular scores and residualising each target score against one of four fibroblast-content specifications: a non-overlapping five-gene fibroblast lineage score (PDGFRA, PDGFRB, LUM, DCN and COL14A1), MCP-counter fibroblasts [24], EPIC CAFs [25], or MCP-counter and EPIC jointly. Patient-level bootstrap confidence intervals used 5,000 resamples. The same framework was applied to FAP13-matrix4, overlap-removed SenMayo-FAP13 and SenMayo-matrix4. Because deconvolution estimates depend on reference composition and preprocessing choices [26], marginal correlations described tissue-level covariation rather than activation independent of abundance.")

    document.add_heading("2.9 Matrix and receptor boundary analyses", level=2)
    add_body(document, "In tumour fibroblasts, full-library log1p(counts per 10,000) expression was used in mixed models of matrix4 or receptor2 on FAP with patient random intercepts. Patient-level fibroblast FAP, matrix4 and SenMayo summaries were correlated by Spearman method. CPTAC protein analyses used a FAP-excluded matrix score (COL1A1, COL1A2 and FN1) to avoid self-correlation. Receptor analyses tested SDC4/CD44 co-induction as a boundary condition; lack of co-induction was not interpreted as evidence against physical receptor function.")

    document.add_heading("2.10 Cognate SASP-immune receptor screen", level=2)
    add_body(document, "Fourteen prespecified, biologically cognate ligand-receptor comparisons were tested across tumour patients by correlating FAP-detected fibroblast ligand expression with T-cell or myeloid receptor expression: IL6-IL6R/IL6ST, CXCL8-CXCR1/CXCR2, CCL2-CCR2 and TGFB1-TGFBR1/TGFBR2. Benjamini-Hochberg correction was applied across all comparisons. Because transcript co-expression is an inference input rather than a direct measurement of signalling [27], these correlations were reported as a receptor screen rather than cell-cell communication.")

    document.add_heading("2.11 Statistical software and reproducibility", level=2)
    add_body(document, "All tests were two-sided. Effect estimates and confidence intervals were prioritized over cell-level P values. R 4.6.1 was invoked explicitly with --vanilla and the R 4.6 user library; R 4.5.2 and 4.6.0 backups were not used. Core packages included SeuratObject, Matrix, lme4, lmerTest, ggplot2 and patchwork. Analyses avoided the R 4.5.2-compiled data.table binary. Figure exports were 600 dpi RGB PNG and LZW-compressed TIFF, with Arial text.")

    document.add_heading("2.12 Data and code availability", level=2)
    add_body(document, "TCGA-COAD/READ data are available through UCSC Xena/GDC; GSE132465, GSE166555 and GSE39582 through GEO; and CPTAC through cBioPortal. Dataset accessions and primary publications are cited above.")
    add_placeholder(document, "SUBMISSION BLOCKER: replace this sentence with the public GitHub repository URL and archived Zenodo DOI containing frozen scripts, derived source tables, checksums and session information.")

    document.add_heading("3. Results", level=1)
    document.add_heading("3.1 Tumour-restricted analysis populations", level=2)
    add_body(document, "GSE132465 contained 1,501 tumour fibroblast-lineage cells and 17,469 tumour epithelial cells from 23 patients (Table 1). The source object also included 1,961 normal fibroblast-lineage and 1,070 normal epithelial cells. FAP detection differed markedly by tissue class: 925 of 1,025 FAP-detected fibroblast-lineage cells were tumour-derived, whereas 1,861 of 2,437 FAP-undetected fibroblast-lineage cells were normal-derived. All compartment and FAP-specific inference therefore used the tumour-restricted population.")

    document.add_heading("3.2 Senescence-associated transcription was enriched in the tumour fibroblast compartment", level=2)
    add_body(document, "Using the frozen 119-gene, FAP-overlap-removed score, tumour fibroblast-lineage cells had higher SenMayo values than paired epithelial cells in all 23 patients (median standardized paired difference 0.466, 95% bootstrap CI 0.410-0.553; exact P = 2.38 x 10^-7; Figure 1A and Table 2). Requiring at least 20 cells in each compartment retained 15 patients and gave the same median difference (0.466, 95% CI 0.410-0.551; P = 6.10 x 10^-5). SASP25 was higher in fibroblasts in 22 of 23 patients (mean log-normalized paired difference 0.236; P = 4.77 x 10^-7; Figure 1B). MKI67 showed the opposite direction in 21 of 23 patients (mean paired difference -0.104; P = 1.81 x 10^-4; Figure 1C). The paired results localize a senescence-associated transcriptional pattern to the fibroblast compartment; stable arrest remains a tissue- and function-level question.")

    document.add_heading("3.3 Expression-matched nulls supported SenMayo specificity but qualified SASP", level=2)
    add_body(document, "The observed mean patient-level SenMayo fibroblast-minus-epithelial standardized difference was 0.499. Across 5,000 expression-matched gene sets, the null median was 0.219 and the 97.5th percentile was 0.340, yielding an empirical upper-tail P = 2.00 x 10^-4 (Figure 3A). For SASP25, the observed difference was 0.739 versus a null median of 0.405 and 97.5th percentile of 0.759 (empirical P = 0.0336; Figure 3B). Many matched signatures were concordantly higher in fibroblasts across patients, showing that concordance alone is not evidence of senescence. The magnitude of SenMayo enrichment, but not SASP to the same degree, exceeded generic lineage-skewed expression.")

    if external:
        sen = external["SenMayo119"]
        sasp = external["SASP25"]
        document.add_heading("3.4 External single-cell evaluation", level=2)
        correlations = external["fibroblast_patient_correlations"]
        add_body(document, f"In GSE166555, {external['eligible_patients']} patients met the frozen cell-count thresholds, and 114 of the 119 frozen SenMayo targets were represented with non-zero variance. SenMayo was higher in fibroblast-like than epithelial pseudobulks in {sen['concordant']} of {sen['n']} patients (median paired difference {sen['median_difference']:.3f}, 95% bootstrap CI {sen['ci_low']:.3f}-{sen['ci_high']:.3f}; exact P = {sen['p_exact']:.3g}; Figure 4). SASP25 was higher in {sasp['concordant']} of {sasp['n']} patients (median paired difference {sasp['median_difference']:.3f}, exact P = {sasp['p_exact']:.3g}), whereas MKI67 was lower in 9 of 10 (median difference {external['MKI67']['median_difference']:.3f}; P = {external['MKI67']['p_exact']:.3g}). Within external fibroblast pseudobulks, FAP13 tracked matrix4 (rho = {correlations['FAP13_matrix4']['rho']:.3f}, P = {correlations['FAP13_matrix4']['p']:.3g}) but not SenMayo (rho = {correlations['FAP13_SenMayo']['rho']:.3f}, P = {correlations['FAP13_SenMayo']['p']:.3g}); SenMayo did not track matrix4 (rho = {correlations['SenMayo_matrix4']['rho']:.3f}, P = {correlations['SenMayo_matrix4']['p']:.3g}). This analysis reproduced the compartment direction and separation of axes in an external cohort; differences in annotation and pseudobulk normalization preclude treating effect sizes as directly interchangeable.")
        next_section = 5
    else:
        next_section = 4

    document.add_heading(f"3.{next_section} FAP-specific effects attenuated after depth and subtype adjustment", level=2)
    add_body(document, "Within the 1,501 tumour fibroblast-lineage cells, the FAP-detected group had greater sequencing depth and higher unadjusted detection of several markers. After adjustment for log UMI depth, fibroblast subtype, MKI67 and patient clustering, the FAP-detection coefficient was 0.0238 for SenMayo (95% CI -0.0051 to 0.0527; P = 0.106) and 0.0323 for SASP25 (95% CI -0.0429 to 0.1074; P = 0.400; Figure 2A). Marker-detection estimates were similarly centred near the null: CDKN2A OR 0.89 (P = 0.529), CDKN2B OR 0.98 (P = 0.911), CDKN1A OR 0.96 (P = 0.788), LMNB1 OR 0.86 (P = 0.601) and MKI67 OR 0.84 (P = 0.614; Figure 2B). The LMNB1 model was singular, and its unadjusted direction was higher in FAP-detected cells. The adjusted results place SenMayo and SASP enrichment at the broader fibroblast-compartment level rather than within the FAP-detected subgroup.")

    document.add_heading(f"3.{next_section + 1} Senescence-associated and FAP-matrix signals were separable within tumour fibroblasts", level=2)
    add_body(document, "At the patient level, fibroblast FAP expression correlated with matrix4 (rho = 0.649, P = 8.03 x 10^-4) but not significantly with SenMayo (rho = 0.331, P = 0.123); SenMayo and matrix4 were themselves unrelated (rho = -0.048, P = 0.826; Figure 2C-E). In full-library mixed models, matrix4 increased with FAP (slope 0.253, 95% CI 0.213-0.294; P = 9.17 x 10^-33), whereas receptor2 did not (slope -0.021, 95% CI -0.062 to 0.019; P = 0.302). These results separate a FAP-associated matrix state from the compartment-level SenMayo pattern and from SDC4/CD44 co-induction.")

    document.add_heading(f"3.{next_section + 2} Bulk covariation was strongly composition dependent", level=2)
    add_body(document, "Marginal FAP13-matrix4 correlations were 0.930 in TCGA-COAD/READ and 0.913 in GSE39582. In TCGA, the partial correlation remained 0.650 with the five-gene lineage proxy but fell to 0.113 with MCP-counter fibroblasts, 0.288 with EPIC CAFs and 0.150 with joint MCP-counter/EPIC adjustment (Figure 3C and Table 3). Thus, the large tissue-level association primarily tracked fibroblast abundance, and the residual estimate depended on the proxy. In contrast, overlap-removed SenMayo-FAP13 partial correlations were 0.378-0.575 across the four specifications. SenMayo-matrix4 estimates ranged from -0.098 to 0.133 and included zero under three of four specifications. Bulk data therefore support composition-adjusted SenMayo-FAP covariation but not an independent senescence-associated matrix axis.")

    document.add_heading(f"3.{next_section + 3} Protein and immune analyses constrained mechanism", level=2)
    immune_figure_number = 5 if external else 4
    add_body(document, f"CPTAC supported FAP covariation with a FAP-excluded COL1A1/COL1A2/FN1 protein score (rho = 0.812, 95% CI 0.706-0.879; FDR = 1.3 x 10^-23), but receptor proteins were not co-induced with the matrix program. This is protein-level evidence for FAP-ECM covariation, not for SASP-ECM coupling. In the tumour-only patient-level cognate screen, none of 14 SASP ligand-immune receptor comparisons survived correction (minimum q = 0.598; Figure {immune_figure_number}). Correlation therefore did not support a specific immune communication axis. Continuous matrix and senescence scores were not independently prognostic in the two bulk cohorts (all Cox P > 0.10), further limiting clinical interpretation.")

    document.add_heading("4. Discussion", level=1)
    document.add_heading("4.1 A reproducible fibroblast-enriched transcriptional pattern", level=2)
    add_body(document, "The analysis identifies a fibroblast-enriched SenMayo transcriptional pattern within CRC tumours. The direction was shared by all 23 discovery patients, retained under a cell-count threshold, exceeded 5,000 expression-matched null signatures and recurred in an external 10-patient pseudobulk evaluation. This convergence makes the compartment distribution reproducible across the available transcriptomic data. SenMayo and SASP nevertheless include inflammatory, growth-factor and stress-response genes that can be expressed without durable arrest, while epithelial-fibroblast comparisons retain lineage identity. The appropriate biological level is therefore a fibroblast-enriched senescence-associated transcriptional state, with cellular senescence reserved for future orthogonal validation.")

    document.add_heading("4.2 FAP marks a matrix state distinct from the compartment SenMayo signal", level=2)
    add_body(document, "FAP detection was strongly related to tumour origin and sequencing depth in the full source object. Restricting analysis to tumour fibroblasts and adjusting for log UMI depth and subtype reduced the SenMayo and SASP coefficients to estimates compatible with no FAP-specific enrichment; individual marker-detection models showed the same pattern. By contrast, FAP retained a strong association with matrix4 at both cell and patient levels. The resulting model contains two related stromal features at different scales: FAP identifies a matrix-linked fibroblast state, while senescence-associated transcription is enriched across the broader fibroblast compartment.")

    document.add_heading("4.3 Relation to recent CRC senescent-fibroblast studies", level=2)
    add_body(document, "The present findings complement two recent mechanistic CRC studies. Ge et al. combined single-cell evidence with multiplex imaging and functional experiments to link senescent fibroblasts with CD8-positive T-cell dysfunction [14]. Yang et al. used computational, spatial, organoid and genetic models to establish macrophage-induced CAF senescence and chemoresistance [15]. Their experimental evidence defines functional senescent-fibroblast phenotypes. Our composition-aware contribution is different: it shows that a patient-paired senescence-associated transcriptional contrast can coexist with a separable FAP-matrix axis, and it specifies where tissue and perturbation experiments are required to connect the two.")

    document.add_heading("4.4 Why bulk correlations overstated a unified senescence-matrix program", level=2)
    add_body(document, "FAP13 and matrix4 were almost perfectly correlated in bulk tissue, but both are sensitive to the fraction of fibroblasts in a specimen. The marked attenuation under MCP-counter and EPIC adjustment shows that the marginal coefficient cannot be interpreted as activation-state coupling independent of abundance. SenMayo-FAP13 covariation was more stable, whereas SenMayo-matrix4 was not. Together with the tumour-fibroblast patient summaries, this pattern argues for partially separable axes: fibroblast abundance/FAP-matrix state and fibroblast-enriched senescence-associated transcription. The five-gene fib5 result alone would have produced an overly favourable and proxy-dependent conclusion.")

    document.add_heading("4.5 Mechanistic and translational limits", level=2)
    add_body(document, "The cognate SASP-receptor screen yielded no multiplicity-supported immune pair, and CPTAC did not measure a coordinated SASP protein program. SDC4/CD44 non-co-induction does not exclude receptor function, but it prevents their presentation as a validated interface. No senolytic, senomorphic or stroma-directed therapeutic inference follows from these data. A clinically meaningful test would require multi-marker tissue localization, direct arrest assays and perturbation, rather than stronger wording around observational scores.")

    document.add_heading("4.6 Limitations and falsifiable next steps", level=2)
    add_body(document, "First, the primary single-cell analysis used one discovery cohort and source annotations; external data differed in platform and cell labels. Second, fibroblast-epithelial score contrasts may retain unmeasured lineage biology despite expression matching. Third, FAP detection is dropout-prone and not equivalent to FAP protein. Fourth, mixed models with 23 patients cannot fully characterize random slopes or patient heterogeneity. Fifth, bulk deconvolution proxies are imperfect and produced materially different residual estimates. Sixth, the analyses were exploratory and not preregistered. Seventh, no spatial co-localization, senescence-associated beta-galactosidase, persistent arrest, DNA-damage response or functional perturbation was available.")
    add_body(document, "The minimum decisive experiment is multiplexed tissue imaging of FAP together with p16 and p21, loss of lamin B1 [10], a proliferation marker and epithelial/stromal segmentation in an independent CRC series, followed by fibroblast isolation or organoid-CAF co-culture to test persistent arrest and SASP secretion. FAP-high and FAP-low fibroblasts should be matched for viability and activation subtype. Concordant tissue localization and persistent, depth-independent protein or functional evidence would determine whether a FAP-specific senescence phenotype exists.")

    document.add_heading("5. Conclusion", level=1)
    add_body(document, "CRC transcriptomes contain a reproducible fibroblast-enriched senescence-associated pattern whose magnitude exceeds expression-matched null expectations and recurs in an external cohort. FAP defines a distinct ECM-rich state, while depth-adjusted tumour-fibroblast analyses place the SenMayo signal at the compartment rather than FAP-detected subgroup level. Together, the two axes provide a focused framework for multi-marker tissue localization and functional testing of stromal senescence in CRC.")

    document.add_heading("Declarations", level=1)
    add_inline_label(document, "Funding: ", "National Natural Science Foundation of China (No. 82360467); Science and Technology Fund Project of Guizhou Provincial Science and Technology Program (No. QKH JC-ZK [2022]-358 and [2022]-348). The funders had no role in study design, data collection and analysis, decision to publish, or preparation of the manuscript.")
    add_inline_label(document, "Competing interests: ", "The authors declare no competing financial or personal interests.")
    add_inline_label(document, "Ethics approval: ", "Only publicly available, de-identified data were analysed; no new human or animal data were collected.")
    add_inline_label(document, "AI-assisted writing disclosure: ", "OpenAI Codex (GPT-5; accessed August 2026) was used for language editing, code review, statistical workflow checking and document-formatting support. It did not generate biological measurements, alter computed analysis outputs, determine the authors' final scientific conclusions or qualify for authorship. The authors independently reviewed, verified and revised all AI-assisted outputs and retain full responsibility for the integrity and accuracy of the work.")
    add_inline_label(document, "Author contributions (CRediT): ", "Jinglin Wang: conceptualization, methodology, writing - original draft. Pan Sun: data curation, validation, writing - review and editing. Lei Xu: investigation, visualization. Yao Wang: formal analysis, software. Wei Zhu: data curation, software. Yan Zhang: data curation. Yao Tang: data curation. Juan Yang: conceptualization, data curation, funding acquisition, project administration, supervision, writing - review and editing. All authors read and approved the submitted version and agree to be accountable for the integrity of the work.")

    document.add_heading("References", level=1)
    references = [
        "Bray Freddie, Laversanne Mathieu, Sung Hyuna, Ferlay Jacques, Siegel Rebecca L., Soerjomataram Isabelle, Jemal Ahmedin. Global cancer statistics 2022: GLOBOCAN estimates of incidence and mortality worldwide for 36 cancers in 185 countries. CA: A Cancer Journal for Clinicians. 2024;74;229-263. doi:10.3322/caac.21834",
        "Sahai Erik, Astsaturov Igor, Cukierman Edna, et al. A framework for advancing our understanding of cancer-associated fibroblasts. Nature Reviews Cancer. 2020;20;174-186. doi:10.1038/s41568-019-0238-1",
        "Fitzgerald Allison A., Weiner Louis M. The role of fibroblast activation protein in health and malignancy. Cancer and Metastasis Reviews. 2020;39;783-803. doi:10.1007/s10555-020-09909-3",
        "Isella Claudio, Terrasi Andrea, Bellomo Sara Erika, et al. Stromal contribution to the colorectal cancer transcriptome. Nature Genetics. 2015;47;312-319. doi:10.1038/ng.3224",
        "Calon Alexandre, Lonardo Enza, Berenguer-Llergo Antonio, et al. Stromal gene expression defines poor-prognosis subtypes in colorectal cancer. Nature Genetics. 2015;47;320-329. doi:10.1038/ng.3225",
        "Qi Jingjing, Sun Hongxiang, Zhang Yao, et al. Single-cell and spatial analysis reveal interaction of FAP+ fibroblasts and SPP1+ macrophages in colorectal cancer. Nature Communications. 2022;13;1742. doi:10.1038/s41467-022-29366-6",
        "Saul Dominik, Kosinsky Robyn Laura, Atkinson Elizabeth J., et al. A new gene set identifies senescent cells and predicts senescence-associated pathways across tissues. Nature Communications. 2022;13;4827. doi:10.1038/s41467-022-32552-1",
        "Gonzalez-Gualda Estela, Baker Andrew G., Fruk Ljiljana, Munoz-Espin Daniel. A guide to assessing cellular senescence in vitro and in vivo. FEBS Journal. 2021;288;56-80. doi:10.1111/febs.15570",
        "Suryadevara Vidyani, Hudgins Alexander D., Rajesh Arjun, et al. SenNet recommendations for detecting senescent cells in different tissues. Nature Reviews Molecular Cell Biology. 2024;25;1001-1023. doi:10.1038/s41580-024-00738-8",
        "Freund Adam, Laberge Remi-Martin, Demaria Marco, Campisi Judith. Lamin B1 loss is a senescence-associated biomarker. Molecular Biology of the Cell. 2012;23;2066-2075. doi:10.1091/mbc.E11-10-0884",
        "Faget Douglas V., Ren Qihao, Stewart Sheila A. Unmasking senescence: context-dependent effects of SASP in cancer. Nature Reviews Cancer. 2019;19;439-453. doi:10.1038/s41568-019-0156-2",
        "Belle Jad I., Sen Devashish, Baer John M., et al. Senescence Defines a Distinct Subset of Myofibroblasts That Orchestrates Immunosuppression in Pancreatic Cancer. Cancer Discovery. 2024;14;1324-1355. doi:10.1158/2159-8290.CD-23-0428",
        "Hattangady Namita Ganesh, Carter Kelly, Maroni-Rana Brett, et al. Mapping the core senescence phenotype of primary human colon fibroblasts. Aging. 2024;16;3068-3087. doi:10.18632/aging.205577",
        "Ge Mengxiao, Sun Shuangyi, Chen Wentong, Xu Zhenxiao, Kang Deyi, Wang Zeqin, Guo Yumeng. Senescent fibroblasts drive CD8+ T cell dysfunction in colorectal cancer via CD36-mediated lipid transfer and peroxidation. Journal of Translational Medicine. 2026;24;310. doi:10.1186/s12967-025-07636-3",
        "Yang Shuaixi, Chen Yukang, Liu Shutong, et al. Macrophage-Induced Senescent Cancer-Associated Fibroblasts Promote SASP-Mediated Chemoresistance in Colorectal Cancer. Cancer Research. 2026; online ahead of print. doi:10.1158/0008-5472.CAN-25-4870",
        "Marisa Laetitia, de Reynies Aurelien, Duval Alex, et al. Gene Expression Classification of Colon Cancer into Molecular Subtypes: Characterization, Validation, and Prognostic Value. PLoS Medicine. 2013;10;e1001453. doi:10.1371/journal.pmed.1001453",
        "Lee Hae-Ock, Hong Yourae, Etlioglu Hakki Emre, et al. Lineage-dependent gene expression programs influence the immune landscape of colorectal cancer. Nature Genetics. 2020;52;594-603. doi:10.1038/s41588-020-0636-z",
        "Uhlitz Florian, Bischoff Philip, Peidli Stefan, et al. Mitogen-activated protein kinase activity drives cell trajectories in colorectal cancer. EMBO Molecular Medicine. 2021;13;e14123. doi:10.15252/emmm.202114123",
        "Vasaikar Suhas, Huang Chen, Wang Xiaojing, et al. Proteogenomic Analysis of Human Colon Cancer Reveals New Therapeutic Opportunities. Cell. 2019;177;1035-1049.e19. doi:10.1016/j.cell.2019.03.030",
        "Cerami Ethan, Gao Jianjiong, Dogrusoz Ugur, et al. The cBio Cancer Genomics Portal: An Open Platform for Exploring Multidimensional Cancer Genomics Data. Cancer Discovery. 2012;2;401-404. doi:10.1158/2159-8290.CD-12-0095",
        "Gao Jianjiong, Aksoy Bulent Arman, Dogrusoz Ugur, et al. Integrative Analysis of Complex Cancer Genomics and Clinical Profiles Using the cBioPortal. Science Signaling. 2013;6;pl1. doi:10.1126/scisignal.2004088",
        "Zimmerman Kip D., Espeland Mark A., Langefeld Carl D. A practical solution to pseudoreplication bias in single-cell studies. Nature Communications. 2021;12;738. doi:10.1038/s41467-021-21038-1",
        "Squair Jordan W., Gautier Matthieu, Kathe Claudia, et al. Confronting false discoveries in single-cell differential expression. Nature Communications. 2021;12;5692. doi:10.1038/s41467-021-25960-2",
        "Becht Etienne, Giraldo Nicolas A., Lacroix Laetitia, et al. Estimating the population abundance of tissue-infiltrating immune and stromal cell populations using gene expression. Genome Biology. 2016;17;218. doi:10.1186/s13059-016-1070-5",
        "Racle Julien, de Jonge Kaat, Baumgaertner Petra, Speiser Daniel E., Gfeller David. Simultaneous enumeration of cancer and immune cell types from bulk tumor gene expression data. eLife. 2017;6;e26476. doi:10.7554/eLife.26476",
        "Avila Cobos Francisco, Alquicira-Hernandez Jose, Powell Joseph E., Mestdagh Pieter, De Preter Katleen. Benchmarking of cell type deconvolution pipelines for transcriptomics data. Nature Communications. 2020;11;5650. doi:10.1038/s41467-020-19015-1",
        "Armingol Erick, Officer Adam, Harismendy Olivier, Lewis Nathan E. Deciphering cell-cell interactions and communication from gene expression. Nature Reviews Genetics. 2021;22;71-88. doi:10.1038/s41576-020-00292-x",
    ]
    for index, reference in enumerate(references, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(-0.65)
        paragraph.paragraph_format.left_indent = Cm(0.65)
        paragraph.add_run(f"{index}. {reference}")

    external_cells = "10 eligible patients" if external else "external analysis pending at manuscript-build time"
    add_table(
        document,
        "Table 1. Audited datasets and analysis populations.",
        ["Dataset", "Level", "Primary analysis population", "Role"],
        [
            ["TCGA-COAD/READ", "Bulk RNA", "380 primary tumours", "Composition-aware discovery"],
            ["GSE39582", "Bulk microarray", "566 tumour samples", "Cross-platform tissue-level replication"],
            ["GSE132465", "Single cell", "63,689 total; tumour-only 1,501 fibroblast-lineage + 17,469 epithelial; 23 patients", "Primary compartment and FAP-specific analyses"],
            ["GSE166555", "Single cell", f"68,702 total; 12 tumour patients; {external_cells}", "External compartment evaluation"],
            ["CPTAC colon", "Proteomics", "97 tumours", "FAP-ECM and receptor boundary"],
        ],
        "The full lineage universe of 3,462 fibroblast-lineage plus 18,539 epithelial cells combined tumour and normal tissue; compartment-level inference was restricted to source-annotated tumour cells. CAF denotes tumour-derived fibroblast-lineage cells."
    )

    add_table(
        document,
        "Table 2. GSE132465 tumour-only senescence-associated analyses.",
        ["Analysis", "Estimate (95% CI where available)", "P", "Interpretation"],
        [
            ["SenMayo119: fibroblast vs epithelial, 23 patients", "Median paired difference 0.466 (0.410-0.553); 23/23 higher", "2.38 x 10^-7", "Robust compartment score difference"],
            ["SenMayo119: >=20 cells/compartment, 15 patients", "Median paired difference 0.466 (0.410-0.551)", "6.10 x 10^-5", "Cell-count sensitivity retained"],
            ["SenMayo119 matched-null specificity", "Observed mean 0.499; null median 0.219; null 97.5% 0.340", "Empirical 2.00 x 10^-4", "Magnitude exceeds matched expression background"],
            ["SASP25: fibroblast vs epithelial", "Mean log-normalized difference 0.236; 22/23 higher", "4.77 x 10^-7", "Directional contrast; weaker signature specificity"],
            ["SASP25 matched-null specificity", "Observed 0.739; null median 0.405; null 97.5% 0.759", "Empirical 0.0336", "Borderline upper-tail specificity"],
            ["MKI67: fibroblast vs epithelial", "Mean difference -0.104; 21/23 lower", "1.81 x 10^-4", "Opposite proliferative distribution"],
            ["Adjusted FAP-detection coefficient: SenMayo119", "beta 0.0238 (-0.0051 to 0.0527)", "0.106", "No independent FAP-specific enrichment"],
            ["Adjusted FAP-detection coefficient: SASP25", "beta 0.0323 (-0.0429 to 0.1074)", "0.400", "No independent FAP-specific enrichment"],
        ],
        "SenMayo119 excludes four FAP13 overlaps and uses the 119 genes represented in GSE132465. FAP mixed models adjusted for log UMI depth, fibroblast subtype, MKI67 and patient clustering."
    )

    add_table(
        document,
        "Table 3. TCGA composition-adjusted partial Spearman correlations.",
        ["Adjustment", "FAP13-matrix4", "SenMayo-FAP13", "SenMayo-matrix4"],
        [
            ["fib5", "0.650 (0.574-0.715)", "0.388 (0.284-0.485)", "0.133 (0.020-0.248)"],
            ["MCP-counter fibroblasts", "0.113 (-0.015 to 0.250)", "0.378 (0.270-0.476)", "-0.098 (-0.221 to 0.043)"],
            ["EPIC CAFs", "0.288 (0.176-0.398)", "0.575 (0.490-0.648)", "0.082 (-0.015 to 0.188)"],
            ["MCP-counter + EPIC", "0.150 (0.040-0.257)", "0.384 (0.273-0.485)", "-0.039 (-0.137 to 0.059)"],
        ],
        "Values are partial Spearman rho with 95% patient-bootstrap confidence intervals. The wide change across fibroblast proxies is itself a sensitivity result and precludes a single proxy-independent activation estimate."
    )

    add_table(
        document,
        "Table 4. Cross-layer evidence and mechanistic boundaries.",
        ["Analysis", "Result", "Evidence boundary"],
        [
            ["TCGA marginal FAP13-matrix4", "rho = 0.930", "Strong tissue covariation; composition sensitive"],
            ["GSE39582 marginal FAP13-matrix4", "rho = 0.913", "Cross-platform tissue-level replication"],
            ["Tumour fibroblast mixed model: matrix4 on FAP", "slope 0.253 (0.213-0.294), P = 9.17 x 10^-33", "FAP-linked matrix state"],
            ["Tumour fibroblast mixed model: receptor2 on FAP", "slope -0.021 (-0.062 to 0.019), P = 0.302", "No receptor co-induction"],
            ["Patient fibroblast FAP vs matrix4", "rho = 0.649, P = 8.03 x 10^-4", "Patient-level matrix coupling"],
            ["Patient fibroblast FAP vs SenMayo", "rho = 0.331, P = 0.123", "No patient-level FAP-senescence confirmation"],
            ["Patient fibroblast SenMayo vs matrix4", "rho = -0.048, P = 0.826", "Axes separable"],
            ["External GSE166555 fibroblast FAP13 vs matrix4", "rho = 0.879, P = 8.14 x 10^-4", "External support for FAP-linked matrix state"],
            ["External GSE166555 fibroblast FAP13 vs SenMayo", "rho = -0.067, P = 0.855", "No external FAP-senescence coupling"],
            ["External GSE166555 fibroblast SenMayo vs matrix4", "rho = -0.406, P = 0.244", "No external senescence-matrix coupling"],
            ["CPTAC FAP vs FAP-excluded matrix protein score", "rho = 0.812 (0.706-0.879), FDR = 1.3 x 10^-23", "Protein FAP-ECM covariation; not SASP evidence"],
            ["Cognate SASP-immune receptor screen", "14 comparisons; minimum q = 0.598", "No multiplicity-supported communication axis"],
        ],
        "rho, Spearman correlation; FDR, Benjamini-Hochberg false-discovery rate. Correlation and co-expression do not establish direction, contact or causality."
    )

    document.add_page_break()
    document.add_heading("Figure legends", level=1)
    legends = [
        "Figure 1. Tumour-only, patient-paired compartment distribution in GSE132465. (A) SenMayo119 score, (B) SASP25 score and (C) MKI67 expression in fibroblast-lineage and epithelial cells. Points are patient means; grey lines connect the same patient; black bars are medians. Exact paired P values are shown. Scores localize transcriptional features and do not diagnose senescence.",
        "Figure 2. FAP-specific inference and separation of fibroblast axes in GSE132465. (A) Adjusted FAP-detection effects on SenMayo119 and SASP25. (B) Adjusted odds ratios for marker detection. (C-E) Patient-level fibroblast correlations among FAP, SenMayo119 and matrix4. Mixed models account for patient clustering; marker models adjust for log UMI depth and subtype. The LMNB1 model was singular.",
        "Figure 3. Signature specificity and bulk composition sensitivity. (A, B) Expression-matched null distributions for the mean patient-level fibroblast-minus-epithelial SenMayo119 and SASP25 contrasts; red lines show observed values. (C) TCGA partial Spearman correlations under four fibroblast-content specifications. Error bars are 95% patient-bootstrap confidence intervals.",
    ]
    if external and EXT_FIGURE.exists():
        legends.append("Figure 4. External tumour-compartment evaluation in GSE166555. Patient-paired pseudobulk SenMayo119, SASP25 and MKI67 values are shown for eligible CAF-like and epithelial compartments. This analysis tests direction in a differently annotated cohort and is not direct evidence of cellular senescence.")
        legends.append("Figure 5. Tumour-only patient-level screen of cognate SASP ligand-receptor pairs in GSE132465. Values are Spearman correlations between FAP-detected fibroblast ligands and T-cell or myeloid receptor expression. No comparison survived Benjamini-Hochberg correction; correlations do not establish communication.")
    else:
        legends.append("Figure 4. Tumour-only patient-level screen of cognate SASP ligand-receptor pairs in GSE132465. Values are Spearman correlations between FAP-detected fibroblast ligands and T-cell or myeloid receptor expression. No comparison survived Benjamini-Hochberg correction; correlations do not establish communication.")
    for legend in legends:
        add_body(document, legend)

    add_figure(document, "Figure 1", FIGURE_DIR / "Figure1_tumor_only_compartment.png")
    add_figure(document, "Figure 2", FIGURE_DIR / "Figure2_FAP_and_matrix_separation.png")
    add_figure(document, "Figure 3", FIGURE_DIR / "Figure3_specificity_and_composition.png")
    if external and EXT_FIGURE.exists():
        add_figure(document, "Figure 4", EXT_FIGURE)
        add_figure(document, "Figure 5", FIGURE_DIR / "Figure6_cognate_SASP_immune_screen.png")
    else:
        add_figure(document, "Figure 4", FIGURE_DIR / "Figure6_cognate_SASP_immune_screen.png")

    path = OUT / "FAP_Senescence_AJCR_rebuilt_draft.docx"
    document.save(path)
    word_count = len(re.findall(r"\b[\w+-]+\b", abstract))
    print(f"Wrote {path}")
    print(f"Abstract word count: {word_count}")
    return path


if __name__ == "__main__":
    build()
