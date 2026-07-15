from __future__ import annotations

import csv
import hashlib
import re
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[3]
FINAL_DIR = ROOT / "work" / "revision_round3" / "final"
REPORT_DIR = ROOT / "work" / "revision_round3" / "reports"
FIGURE_DIR = ROOT / "outputs" / "figures"

DOCUMENTS = {
    "manuscript": FINAL_DIR / "JCMM_Manuscript_Revised_Round3.docx",
    "supplement": FINAL_DIR / "JCMM_Supplementary_Material_Round3.docx",
    "review": FINAL_DIR / "JCMM_Scientific_Review_Report_Round3.docx",
    "checklist": FINAL_DIR / "JCMM_Submission_Materials_Checklist_Round3.docx",
    "cover_letter": FINAL_DIR / "JCMM_Cover_Letter_Round3.docx",
}


def sha256_bytes(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def media_hashes(docx_path: Path) -> list[str]:
    with zipfile.ZipFile(docx_path) as archive:
        names = sorted(
            name for name in archive.namelist() if name.startswith("word/media/")
        )
        return [sha256_bytes(archive.read(name)) for name in names]


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def inspect_document(label: str, path: Path) -> dict[str, object]:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    hashes = media_hashes(path)
    markdown_tokens = sorted(set(re.findall(r"(?:\*\*|^#{1,4}\s|`)", text, re.M)))
    return {
        "document": label,
        "path": str(path.relative_to(ROOT)),
        "size_bytes": path.stat().st_size,
        "sha256": file_hash(path),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "media_files": len(hashes),
        "replacement_characters": text.count("\ufffd"),
        "markdown_tokens": ";".join(markdown_tokens),
        "repository_placeholder": text.count("repository URL and Zenodo DOI pending"),
        "trailing_empty_paragraph": int(
            bool(document.paragraphs) and not document.paragraphs[-1].text.strip()
        ),
        "text_characters": len(text),
    }


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    rows = [inspect_document(label, path) for label, path in DOCUMENTS.items()]

    manuscript_expected = {
        file_hash(FIGURE_DIR / "main" / f"Figure_{number:02d}.png")
        for number in range(1, 7)
    }
    supplement_expected = {
        file_hash(FIGURE_DIR / "supplementary" / f"Figure_S{number:02d}.png")
        for number in range(1, 6)
    }
    manuscript_embedded = set(media_hashes(DOCUMENTS["manuscript"]))
    supplement_embedded = set(media_hashes(DOCUMENTS["supplement"]))

    checks = {
        "manuscript_has_six_current_figures": manuscript_embedded == manuscript_expected,
        "supplement_has_five_current_figures": supplement_embedded == supplement_expected,
        "manuscript_placeholder_present_once": next(
            row["repository_placeholder"] for row in rows if row["document"] == "manuscript"
        )
        == 1,
        "no_replacement_characters": all(
            row["replacement_characters"] == 0 for row in rows
        ),
        "no_markdown_tokens": all(not row["markdown_tokens"] for row in rows),
        "checklist_has_no_trailing_empty_paragraph": next(
            row["trailing_empty_paragraph"] for row in rows if row["document"] == "checklist"
        )
        == 0,
    }

    csv_path = REPORT_DIR / "document_structure_audit.csv"
    with csv_path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)

    md_path = REPORT_DIR / "document_structure_audit.md"
    with md_path.open("w", encoding="utf-8") as handle:
        handle.write("# Final DOCX structural audit\n\n")
        handle.write("| Document | Paragraphs | Tables | Images | Placeholder | SHA-256 |\n")
        handle.write("|---|---:|---:|---:|---:|---|\n")
        for row in rows:
            handle.write(
                f"| {row['document']} | {row['paragraphs']} | {row['tables']} | "
                f"{row['media_files']} | {row['repository_placeholder']} | "
                f"`{row['sha256']}` |\n"
            )
        handle.write("\n## Checks\n\n")
        for name, passed in checks.items():
            handle.write(f"- {'PASS' if passed else 'FAIL'}: {name}\n")

    failed = [name for name, passed in checks.items() if not passed]
    print(csv_path.relative_to(ROOT))
    print(md_path.relative_to(ROOT))
    print(f"checks={len(checks)} failed={len(failed)}")
    if failed:
        raise SystemExit("Failed checks: " + ", ".join(failed))


if __name__ == "__main__":
    main()
