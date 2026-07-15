from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
FINAL_DIR = ROOT / "work" / "revision_round3" / "final"
FIGURE_DIR = ROOT / "outputs" / "figures"

MANUSCRIPT_MD = FINAL_DIR / "JCMM_Manuscript_Revised_Round3.md"
REVIEW_MD = FINAL_DIR / "JCMM_Scientific_Review_Report_Round3.md"
CHECKLIST_MD = FINAL_DIR / "JCMM_Submission_Materials_Checklist_Round3.md"

MANUSCRIPT_DOCX = FINAL_DIR / "JCMM_Manuscript_Revised_Round3.docx"
REVIEW_DOCX = FINAL_DIR / "JCMM_Scientific_Review_Report_Round3.docx"
CHECKLIST_DOCX = FINAL_DIR / "JCMM_Submission_Materials_Checklist_Round3.docx"
SUPPLEMENT_DOCX = FINAL_DIR / "JCMM_Supplementary_Material_Round3.docx"


def set_run_font(run, western: str, east_asia: str, size: float | None = None) -> None:
    run.font.name = western
    if size is not None:
        run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:ascii"), western)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), western)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document, *, language: str, header_text: str = "") -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    if language == "zh":
        western, east_asia, body_size = "Arial", "SimSun", 10.5
        heading_font, heading_east = "Arial", "Microsoft YaHei"
    else:
        western, east_asia, body_size = "Times New Roman", "Times New Roman", 11
        heading_font, heading_east = "Arial", "Arial"

    normal = doc.styles["Normal"]
    normal.font.name = western
    normal.font.size = Pt(body_size)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    for style_name, size, before, after in (
        ("Title", 16, 0, 12),
        ("Heading 1", 14, 12, 6),
        ("Heading 2", 12, 10, 5),
        ("Heading 3", 11, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = heading_font
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 45, 61)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), heading_east)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Reference" not in doc.styles:
        ref_style = doc.styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref_style = doc.styles["Reference"]
    ref_style.base_style = doc.styles["Normal"]
    ref_style.paragraph_format.left_indent = Cm(0.75)
    ref_style.paragraph_format.first_line_indent = Cm(-0.75)
    ref_style.paragraph_format.line_spacing = 1.0
    ref_style.paragraph_format.space_after = Pt(3)

    if "Figure Legend" not in doc.styles:
        legend_style = doc.styles.add_style("Figure Legend", WD_STYLE_TYPE.PARAGRAPH)
    else:
        legend_style = doc.styles["Figure Legend"]
    legend_style.base_style = doc.styles["Normal"]
    legend_style.font.size = Pt(9.5)
    legend_style.paragraph_format.line_spacing = 1.0
    legend_style.paragraph_format.space_after = Pt(7)
    legend_style.paragraph_format.keep_with_next = False

    if header_text:
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_p.add_run(header_text)
        set_run_font(header_run, western, east_asia, 8.5)
        header_run.font.color.rgb = RGBColor(90, 90, 90)
    add_page_number(section.footer.paragraphs[0])

    doc.core_properties.author = "Jinglin Wang et al."
    doc.core_properties.subject = "FAP-CLDN colorectal cancer manuscript revision"
    doc.core_properties.keywords = "FAP; CAF; claudin; colorectal cancer; spatial transcriptomics"


INLINE_TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline_runs(paragraph, text: str, *, language: str, red_placeholder: bool = False) -> None:
    western = "Arial" if language == "zh" else "Times New Roman"
    east_asia = "SimSun" if language == "zh" else "Times New Roman"
    text = text.replace("\\*", "*")
    pos = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, western, east_asia)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_run_font(run, western, east_asia)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Courier New", "Microsoft YaHei", 9)
            run.font.color.rgb = RGBColor(38, 82, 115)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, western, east_asia)

    if red_placeholder or "pending]" in text.lower():
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(192, 0, 0)
            run.bold = True


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_table(doc: Document, rows: list[list[str]], *, language: str) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    total_twips = 9300
    equal_width = total_twips // columns

    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        if row_index == 0:
            set_repeat_table_header(row)
        for col_index in range(columns):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_width(cell, equal_width)
            value = values[col_index] if col_index < len(values) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, value, language=language)
            for run in p.runs:
                run.font.size = Pt(8.5 if language == "en" else 8)
                if row_index == 0:
                    run.bold = True
            if row_index == 0:
                set_cell_shading(cell, "D9EAF2")
    doc.add_paragraph()


def add_markdown_to_doc(
    doc: Document,
    md_path: Path,
    *,
    language: str,
    manuscript: bool = False,
) -> dict[str, str]:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    index = 0
    in_references = False
    figure_legends: dict[str, str] = {}

    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_markdown_table(lines, index)
            add_table(doc, rows, language=language)
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "B7C9D6")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            index += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            hashes, heading_text = heading_match.groups()
            if manuscript and heading_text == "Abstract":
                doc.add_page_break()
            if manuscript and heading_text == "References":
                doc.add_page_break()
                in_references = True
            if len(hashes) == 1:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_runs(p, heading_text, language=language)
            else:
                style = f"Heading {min(len(hashes) - 1, 3)}"
                p = doc.add_paragraph(style=style)
                add_inline_runs(p, heading_text, language=language)
            index += 1
            continue

        if stripped.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, "☐ " + stripped[6:], language=language)
            index += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:], language=language)
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered and not in_references:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, numbered.group(2), language=language)
            index += 1
            continue

        if manuscript and in_references and re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="Reference")
            add_inline_runs(p, stripped, language=language)
            index += 1
            continue

        style = "Normal"
        if manuscript and re.match(r"^\*\*(?:Supplementary )?Figure\s", stripped):
            style = "Figure Legend"
            key_match = re.match(r"^\*\*((?:Supplementary )?Figure\s+[S0-9]+)\.", stripped)
            if key_match:
                figure_legends[key_match.group(1)] = stripped

        p = doc.add_paragraph(style=style)
        if manuscript and stripped.startswith("**Running title:"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if manuscript and stripped.startswith("**Jinglin Wang"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if manuscript and index < 20 and (
            stripped.startswith("1 Department")
            or stripped.startswith("2 Department")
            or stripped.startswith("\\*Co-corresponding")
            or stripped.startswith("Jinglin Wang, Email")
            or stripped.startswith("Juan Yang")
            or stripped.startswith("ORCID")
        ):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(
            p,
            stripped,
            language=language,
            red_placeholder="repository URL and Zenodo DOI pending" in stripped,
        )
        index += 1

    return figure_legends


def add_figure(doc: Document, path: Path, legend: str, *, alt_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.45))
    shape._inline.docPr.set("descr", alt_text)
    legend_p = doc.add_paragraph(style="Figure Legend")
    add_inline_runs(legend_p, legend, language="en")


def build_manuscript() -> dict[str, str]:
    doc = Document()
    configure_document(
        doc,
        language="en",
        header_text="FAP-associated ECM and epithelial claudins in CRC",
    )
    legends = add_markdown_to_doc(doc, MANUSCRIPT_MD, language="en", manuscript=True)
    doc.add_page_break()
    heading = doc.add_paragraph(style="Heading 1")
    add_inline_runs(heading, "Main figures", language="en")
    for number in range(1, 7):
        if number > 1:
            doc.add_page_break()
        key = f"Figure {number}"
        image_path = FIGURE_DIR / "main" / f"Figure_{number:02d}.png"
        add_figure(doc, image_path, legends[key], alt_text=f"{key} of the FAP-CLDN study")
    doc.save(MANUSCRIPT_DOCX)
    return legends


def build_supplement(legends: dict[str, str]) -> None:
    doc = Document()
    configure_document(
        doc,
        language="en",
        header_text="Supplementary material: FAP-associated ECM and epithelial claudins",
    )
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(p, "Supplementary material", language="en")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(
        p,
        "Multiscale analyses distinguish a FAP-associated extracellular matrix programme from epithelial claudin expression in colorectal cancer",
        language="en",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(p, "Jinglin Wang et al.", language="en")
    for number in range(1, 6):
        doc.add_page_break()
        key = f"Supplementary Figure S{number}"
        image_path = FIGURE_DIR / "supplementary" / f"Figure_S{number:02d}.png"
        add_figure(doc, image_path, legends[key], alt_text=key)
    doc.save(SUPPLEMENT_DOCX)


def build_review() -> None:
    doc = Document()
    configure_document(doc, language="zh", header_text="FAP-CLDN第三轮科学审查报告")
    add_markdown_to_doc(doc, REVIEW_MD, language="zh")
    doc.save(REVIEW_DOCX)


def build_checklist() -> None:
    doc = Document()
    configure_document(doc, language="zh", header_text="JCMM投稿材料清单")
    add_markdown_to_doc(doc, CHECKLIST_MD, language="zh")
    if doc.paragraphs and not doc.paragraphs[-1].text.strip():
        paragraph = doc.paragraphs[-1]._element
        paragraph.getparent().remove(paragraph)
    doc.save(CHECKLIST_DOCX)


def main() -> None:
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    legends = build_manuscript()
    build_supplement(legends)
    build_review()
    build_checklist()
    for path in (MANUSCRIPT_DOCX, SUPPLEMENT_DOCX, REVIEW_DOCX, CHECKLIST_DOCX):
        print(f"{path.relative_to(ROOT)}\t{path.stat().st_size}")


if __name__ == "__main__":
    main()
