# -*- coding: utf-8 -*-
"""v4.5 × codex → v5.0 融合：注入三件武器（经验零分布/ABSOLUTE 纯度/CMS 交互）+ 统一 CPTAC 口径 + tissue-state 措辞 + 新图"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = r'C:/Users/Spica/WorkBuddy/Claw/deliverables/FAP_SDC4_CD44_v5.0'
FIGDIR = os.path.join(BASE, '图片源')
OUT = os.path.join(BASE, '论文正文', 'FAP_SDC4_CD44_论文_v5.0_2026-08-06.docx')
os.makedirs(os.path.dirname(OUT), exist_ok=True)

SRC = r'C:/Users/Spica/WorkBuddy/Claw/deliverables/FAP_SDC4_CD44_v4.5/v45_modified_text.txt'
SRC44 = r'C:/Users/Spica/WorkBuddy/Claw/deliverables/fap_early_crc/v44_extract.txt'

# ---------- 1. 读取 v4.5 段落 ----------
with open(SRC, encoding='utf-8') as f:
    content = f.read()
paras = []  # (seq, idx, text)
for line in content.split('\n'):
    m = re.match(r'^\[(\d+)\] (.*)$', line)
    if m:
        paras.append((m.group(1), m.group(2)))

# 表格从 v44_extract 提取（v4.5 已过滤病理行，此处复用）
with open(SRC44, encoding='utf-8') as f:
    c44 = f.read()
if '\n=== TABLES ===\n' in c44:
    tables_part = c44.split('\n=== TABLES ===\n')[1]
else:
    tables_part = ''

# ---------- 2. 定义修改与插入 ----------
# 2.1 替换字典：idx -> (old_substr, new_text) 或 idx -> 全文替换
REPLACE = {}
# Abstract Methods：补三件武器
old_am = "Analyses were performed between July and August 2026"
new_am = None  # 实际改在 Abstract Methods 段 [9] 尾部加句
REPLACE['9'] = [
    ("To address score circularity, primary analyses used non-overlapping scores (FAP13, matrix4, receptor2).",
     "To address score circularity, primary analyses used non-overlapping scores (FAP13, matrix4, receptor2). Robustness was tested with three matched empirical null distributions (10,000 draws each, matched on expression, variability and detection), ABSOLUTE purity adjustment, and CMS stratification with classifier-overlap sensitivity."),
]
# Abstract Results：加零分布/纯度
REPLACE['10'] = [
    ("(8) A transcriptomic stromal signature (80 genes), but not a 6-gene ECM proxy, was independently associated with lymph node metastasis in TCGA after T-stage adjustment (OR = 3.37, P = 0.04); a nomogram achieved AUC = 0.888.",
     "(8) The TCGA FAP13–matrix4 association (rho = 0.930 in 380 barcode-reviewed primary tumours) exceeded three matched empirical null distributions (two-sided empirical P ≤ 9.0 × 10⁻⁴) and survived ABSOLUTE purity adjustment (rho = 0.906) and CMS stratification (within-CMS rho = 0.745–0.945). (9) A transcriptomic stromal signature (80 genes), but not a 6-gene ECM proxy, was independently associated with lymph node metastasis in TCGA after T-stage adjustment (OR = 3.37, P = 0.04); a nomogram achieved AUC = 0.888."),
]
# Methods 2.9 CMS：补交互检验描述
REPLACE['46'] = [
    ("Consensus molecular subtype (CMS) labels were recomputed in 383 tumors with the single-sample predictor (CMSclassifier v1.0.0) [25].",
     "Consensus molecular subtype (CMS) labels were recomputed in 383 tumors with CMSclassifier v1.0.0 [25] using both random-forest (RF; posterior ≥ 0.5) and single-sample predictor (SSP; correlation ≥ 0.15, delta ≥ 0.06) methods. Because 11 RF and 10 SSP classifier input genes overlapped FAP13 or matrix4, both classifiers were re-run after omitting those genes. Within-CMS correlations used 5,000 bootstrap resamples with BH correction within classifier; global interaction models used ranked matrix4 as outcome, ranked FAP13, CMS, project, and the FAP13-by-CMS interaction, with significance estimated by 20,000 project-blocked Freedman–Lane permutations and BH correction across the four full/omitted RF/SSP tests [37]."),
]
# Methods 2.4：补经验零分布 + ABSOLUTE 纯度
REPLACE['34'] = [
    ("Correlations were computed with Spearman's rank method (exact = FALSE).",
     "Correlations were computed with Spearman's rank method (exact = FALSE). To test whether the primary FAP13–matrix4 association exceeded transcriptomic background, three matched empirical null distributions were constructed with 10,000 draws each: random FAP13-like versus random matrix4-like scores, fixed FAP13 versus random matrix4-like scores, and random FAP13-like versus fixed matrix4. For each target gene, 500 candidate genes were selected by proximity in mean expression, log standard deviation and detection rate across 380 barcode-reviewed TCGA tumours; genes were sampled without replacement within each draw, and two-sided empirical P values used (1 + number of absolute null correlations at least as large as observed)/(10,000 + 1). For purity sensitivity, FAP13, matrix4 and ABSOLUTE purity were rank-transformed; FAP13 and matrix4 ranks were residualised separately against purity rank and project, and their residual correlation was bootstrapped (5,000 resamples) [36]."),
]
# Results 3.6 CPTAC：统一口径（0.891 含 FAP vs 0.812 FAP-free matrix3）
REPLACE['76'] = [
    ("FAP–ECM protein score ρ = 0.891, P < 10⁻³³; Table 5).",
     "FAP–ECM protein score ρ = 0.891, P < 10⁻³³; Table 5). A stricter, FAP-excluded matrix3 protein score (COL1A1, COL1A2, FN1) reproduced the association in an independent CPTAC analysis (ρ = 0.812, 95% CI 0.706–0.879; FDR = 1.3 × 10⁻²³; Figure 9A), confirming that the protein-level covariation does not depend on including FAP itself within the matrix score."),
]
# Results 3.1 CMS 段：补 within-CMS 结果 + 交互检验（替换原 CMS4 背景段）
REPLACE['59'] = [
    ("The CMS single-sample predictor assigned 314 of 383 tumors; FAP-CAF scores differed across subtypes (Kruskal–Wallis P = 1.51×10⁻²⁸) and were highest in CMS4 (median 1.126; Figure 1d). The FAP–CMS4 association is consistent with previous CRC studies [6, 28] and is presented as context rather than a new subtype marker; FAP is a component of the CMS4 stromal signature, introducing a definitional circularity that we acknowledge.",
     "The CMS single-sample predictor assigned 314 of 383 tumors; FAP-CAF scores differed across subtypes (Kruskal–Wallis P = 1.51×10⁻²⁸) and were highest in CMS4 (median 1.126; Figure 1d). The FAP–CMS4 association is consistent with previous CRC studies [6, 28] and is presented as context rather than a new subtype marker; FAP is a component of the CMS4 stromal signature, introducing a definitional circularity that we acknowledge. To test whether the FAP13–matrix4 association was confined to CMS4, within-CMS correlations were computed with both RF and SSP classifiers: estimates ranged from 0.745 to 0.945 (full classifiers) and from 0.787 to 0.942 after classifier-overlap gene omission (Supplementary Figure S7). Global interaction tests were non-significant after BH correction across the four full/omitted RF/SSP families (raw RF interaction P = 0.042; FDR = 0.168), indicating that matrix covariation is present across CMS classes rather than being a CMS4-specific phenomenon."),
]
# Discussion 4.2 开头：tissue-state 措辞校准（与 codex 统一）
REPLACE['90'] = [
    ("The convergence of four observations—bulk FAP13–receptor2 null correlations,",
     "The convergence of five observations—bulk FAP13–receptor2 null correlations, matched-null and purity-adjusted robustness of the matrix association (Section 3.10),"),
]
# Limitations：补纯度/零分布口径局限
REPLACE['100'] = [
    ("Eighth, the patient-level single-cell cross-compartment analyses were limited to 15 patients",
     "Eighth, the ABSOLUTE purity-adjusted estimate (ρ = 0.906) adjusts for one global copy-number-based purity proxy and does not resolve detailed cellular composition; matched-null analyses are designed to test transcriptomic background but not all possible biological confounders. Ninth, the patient-level single-cell cross-compartment analyses were limited to 15 patients"),
]
# 标题去 "program" 措辞 → tissue state（正文保持 program/interface 叙事，标题统一）
REPLACE['56'] = [
    ("3.1 Non-overlapping scoring shows that matrix coupling, not receptor co-induction, is intrinsic to the FAP-marked stromal program",
     "3.1 Non-overlapping scoring shows that matrix coupling, not receptor co-induction, is intrinsic to the FAP-marked stromal tissue state"),
]

# 2.2 插入段落（idx -> 新段落列表，插在该 idx 之后）
INSERTS = {
    # 3.10 新小节：插在 3.9 内容 [84] 之后、4. Discussion [85] 之前 → 用 key '84'
    '84': [
        ('3.10', '3.10 The FAP13–matrix4 association exceeds matched transcriptomic backgrounds and is robust to purity and CMS stratification'),
        ('3.10a', 'In 380 barcode-reviewed TCGA-COAD/READ primary tumours, the FAP13–matrix4 correlation was 0.930 (95% CI 0.910–0.943; FDR = 1.3 × 10⁻¹⁶⁵), reproduced on a different platform in GSE39582 (ρ = 0.913, 95% CI 0.895–0.928; FDR = 5.8 × 10⁻²²²). Three matched empirical null distributions (10,000 draws each; candidate genes matched on mean expression, log standard deviation and detection rate) had 97.5th percentiles of 0.703 (random-versus-random), 0.839 (fixed FAP13 versus random matrix4) and 0.832 (random FAP13 versus fixed matrix4); the observed 0.930 exceeded all three distributions (two-sided empirical P from 1.0 × 10⁻⁴ to 9.0 × 10⁻⁴; Figure 9B–D). The association was thus not reproduced by gene sets matched only on abundance, variability and detection.'),
        ('3.10b', 'The association was robust to anatomical site (COAD ρ = 0.939, 95% CI 0.919–0.953, n = 285; READ ρ = 0.899, 95% CI 0.836–0.935, n = 94), to sequential gene and sample omission, and to adjustment for ABSOLUTE purity and project (rank-residual ρ = 0.906, 95% CI 0.879–0.924, n = 351; Figure 9A). These sensitivity analyses indicate that the FAP13–matrix4 covariation reflects a coordinated stromal matrix state rather than score overlap, a single anatomical site, or variable tumour purity.'),
        ('3.10c', 'Figure 9. Matched-null and purity robustness of the FAP13–matrix4 association. (A) CPTAC FAP-excluded matrix3 protein correlation (left) and TCGA rank-residual estimate after ABSOLUTE purity and project adjustment (right). (B–D) Three 10,000-draw matched empirical null distributions (random-versus-random; fixed FAP13 versus random matrix4; random FAP13 versus fixed matrix4); the vertical line marks the observed TCGA ρ = 0.930, and two-sided empirical P values are shown.'),
    ],
    # References 区：在 [35] Yoshihara（idx '153'）后追加 ABSOLUTE [36] + Freedman-Lane [37]
    '153': [
        ('36', '36. Carter SL, Cibulskis K, Helman E, et al. Absolute quantification of somatic DNA alterations in human cancer. Nat Biotechnol. 2012;30(5):413-421. doi:10.1038/nbt.2203'),
        ('37', '37. Freedman D, Lane D. A nonstochastic interpretation of reported significance levels. J Bus Econ Stat. 1983;1(4):292-298. doi:10.1080/07350015.1983.10509368'),
    ],
    # 图注区：在 S6 图注 [190] 后追加 S7 图注（补充图 S7 = codex S1 CMS）
    '190': [
        ('S7cap', 'Supplementary Figure S7. CMS-stratified sensitivity analysis. Within-CMS FAP13–matrix4 correlations from RF and SSP high-confidence calls, before and after omission of classifier input genes overlapping FAP13/matrix4.'),
    ],
}

# ---------- 3. 应用修改 ----------
final = []
for idx, text in paras:
    # 替换
    if idx in REPLACE:
        for old, new in REPLACE[idx]:
            if old in text:
                text = text.replace(old, new)
            else:
                print(f'[WARN] REPLACE miss [{idx}]: {old[:50]}...')
    final.append((idx, text))
    # 插入
    if idx in INSERTS:
        for new_idx, new_text in INSERTS[idx]:
            final.append((new_idx, new_text))

# ---------- 4. 输出修改后文本 ----------
out_txt = os.path.join(BASE, 'v50_modified_text.txt')
with open(out_txt, 'w', encoding='utf-8') as f:
    for idx, text in final:
        f.write(f'[{idx}] {text}\n')
    f.write('\n=== TABLES ===\n')
    f.write(tables_part)
print('v50 modified paragraphs:', len(final))

# 残留/新增检查
full = '\n'.join(t for _, t in final)
for kw in ['empirical null', 'ABSOLUTE', 'Freedman', '0.930', '0.906', '0.745', 'matrix3', 'Supplementary Figure S7', 'Figure 9']:
    print(f'含 {kw!r}: {kw in full}')

# ---------- 5. 重建 docx ----------
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(2.54)
sec.left_margin = sec.right_margin = Cm(2.54)

def set_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_para(text, bold=False, size=12, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    parts = re.split(r'(\[\d+(?:,\s*\d+)*\])', text)
    for part in parts:
        if not part:
            continue
        if re.fullmatch(r'\[\d+(?:,\s*\d+)*\]', part):
            r = p.add_run(part)
            set_font(r, size=size, bold=bold)
            r.font.superscript = True
        else:
            r = p.add_run(part)
            set_font(r, size=size, bold=bold)
    return p

HEADING1 = {'1. Introduction', '2. Materials and Methods', '3. Results', '4. Discussion', '5. Conclusion',
            'Abstract', 'Keywords:', 'References', 'Tables', 'Figures and Figure Legends',
            'Supplementary Figures', 'Declarations', 'Funding:', 'Competing interests:',
            'Ethics approval:', 'AI-assisted writing disclosure:', 'Author contributions (CRediT):'}

# 图映射（v4.5 20 张 + Figure 9 零分布图 + S7 CMS 图）
FIG_MAP = {
    'Figures and Figure Legends': ['image1.png'],
    'Figure 1. Bulk': ['image2.png'],
    'Figure 2. Single-cell': ['image3.png'],
    'Figure 3. CellChat': ['image4.png'],
    'Figure 4. Spatially': ['image5.png'],
    'Figure 5. ECM-SDC4/CD44': ['image6.png', 'image7.png'],
    'Figure 6. Nomogram': ['image8.png', 'image9.png'],
    'Figure 7. MMR-status': ['image10.png'],
    'Figure 9. Matched-null': ['codex_Fig2_robustness_nulls.png'],
    'Supplementary Figure S1.': ['image11.png', 'image12.png'],
    'Supplementary Figure S2.': ['image13.png', 'image14.png'],
    'Supplementary Figure S3.': ['image15.png'],
    'Supplementary Figure S4.': ['image16.png', 'image17.png'],
    'Supplementary Figure S5.': ['image18.png'],
    'Supplementary Figure S6.': ['image19.png', 'image20.png'],
    'Supplementary Figure S7.': ['codex_S1_CMS.png'],
}
flushed = set()

def flush_figs(prefix):
    for key, imgs in FIG_MAP.items():
        if prefix.startswith(key) and key not in flushed:
            flushed.add(key)
            for img in imgs:
                fp = os.path.join(FIGDIR, img)
                if not os.path.exists(fp):
                    print('MISSING:', fp)
                    continue
                from PIL import Image
                im = Image.open(fp)
                w = 14.0
                h = 14.0 * im.size[1] / im.size[0]
                if h > 19:
                    h = 19
                    w = 19 * im.size[0] / im.size[1]
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run()
                r.add_picture(fp, width=Cm(w), height=Cm(h))
            return

# 解析表格（v4.5 表格区，追加 codex 摘要表作为 Table 8）
table_blocks = []
cur = []
for line in tables_part.split('\n'):
    if line.startswith('--- Table '):
        if cur:
            table_blocks.append(cur)
        cur = [line]
    else:
        cur.append(line)
if cur:
    table_blocks.append(cur)
parsed_tables = []
for tb in table_blocks:
    rows = [line.split(' | ') for line in tb[1:] if line.strip()]
    parsed_tables.append({'header': tb[0], 'rows': rows})

# 遍历段落构建
for idx, text in final:
    strip = text.strip()
    # 新参考文献条目（字符串 idx '36'/'37'）→ References 样式
    if idx in ('36', '37'):
        add_para(strip, size=11, space_after=3)
        continue
    try:
        n = int(idx)
    except ValueError:
        n = 999
    strip = text.strip()
    if n <= 104 or (isinstance(idx, str) and idx.startswith('3.10')):
        is_heading = any(strip.startswith(h) for h in HEADING1) or \
                     re.match(r'^\d+\.\d+\s', strip) or re.match(r'^\d+\.\s', strip)
        if is_heading:
            add_para(strip, bold=True, space_after=8)
        else:
            add_para(strip)
    elif n == 105:  # Declarations 边界之后由原逻辑处理
        pass
    elif 105 <= n <= 116:
        if strip.endswith(':'):
            add_para(strip, bold=True, space_after=4)
        else:
            add_para(strip)
    elif n == 117:
        add_para(strip, bold=True, space_after=8)
    elif 118 <= n <= 153:
        add_para(strip, size=11, space_after=3)
    elif n == 154:
        add_para(strip, bold=True, space_after=8)
    elif 155 <= n <= 168:
        if strip.startswith('Table ') or strip.startswith('Note:'):
            add_para(strip, bold=strip.startswith('Table '), size=11, space_after=4)
        else:
            add_para(strip, size=11)
    elif n >= 169 or (isinstance(idx, str) and idx.startswith('S7')):
        flush_figs(strip)
        if strip.startswith(('Figure ', 'Supplementary Figure')):
            add_para(strip, size=11, space_after=10)
        else:
            add_para(strip, size=11)

# 表格重建
for tb in parsed_tables:
    header_line = tb['header']
    m = re.match(r'--- Table (\d+) \((\d+)x(\d+)\) ---', header_line)
    if not m:
        continue
    rows = tb['rows']
    if not rows:
        continue
    n_cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=n_cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = t.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(row[j] if j < len(row) else '')
            set_font(r, size=9, bold=(i == 0))

doc.save(OUT)
print('saved:', OUT)
print('size MB:', round(os.path.getsize(OUT)/1e6, 2))
