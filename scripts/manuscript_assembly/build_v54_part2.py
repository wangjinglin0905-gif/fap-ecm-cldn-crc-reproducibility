# -*- coding: utf-8 -*-
"""v54 part2: docx + md 重建（复用 v52 构建逻辑，基线=v54_modified_text.txt）"""
import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
FIGDIR = os.path.join(BASE, '图片源')
OUT_DOCX = os.path.join(BASE, '论文正文', 'FAP_SDC4_CD44_论文_v5.4_2026-08-08.docx')
OUT_MD = os.path.join(BASE, '论文正文', 'FAP_SDC4_CD44_论文_v5.4_2026-08-08.md')
SRC54 = os.path.join(BASE, 'v54_modified_text.txt')

# ---------- 读取 v54 段落 ----------
paras = []
with open(SRC54, encoding='utf-8') as f:
    for line in f:
        line = line.rstrip('\n').rstrip('\r')
        if not line.strip():
            continue
        m = re.match(r'^\[([^\]]+)\]\s?(.*)$', line)
        if m:
            paras.append((m.group(1), m.group(2)))

# 拆分：表格数据从 v50 的 === TABLES === 段解析（v52 表格未变）；body 用 v54 全部段落
SRC50 = os.path.join(BASE, 'v50_modified_text.txt')
v50_content = open(SRC50, encoding='utf-8').read()
if '=== TABLES ===' in v50_content:
    tables_part = v50_content.split('=== TABLES ===')[1]
else:
    tables_part = ''
body_part = paras

# ---------- docx 重建 ----------
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(2.54)
sec.left_margin = sec.right_margin = Cm(2.54)

def set_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_para(text, bold=False, size=12, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    parts = re.split(r'(\[\d+(?:,\s*\d+)*\])', text)
    for part in parts:
        if not part:
            continue
        if re.fullmatch(r'\[\d+(?:,\s*\d+)*\]', part):
            r = p.add_run(part)
            set_font(r, size=size, bold=bold)
            r.font.superscript = True
        else:
            r = p.add_run(part)
            set_font(r, size=size, bold=bold)
    return p

HEADING1 = {'1. Introduction', '2. Materials and Methods', '3. Results', '4. Discussion', '5. Conclusion',
            'Abstract', 'Keywords:', 'References', 'Tables', 'Figures and Figure Legends',
            'Supplementary Figures', 'Declarations', 'Funding:', 'Competing interests:',
            'Ethics approval:', 'AI-assisted writing disclosure:', 'Author contributions (CRediT):'}

FIG_MAP = {
    # 修正 2026-08-08：主图 Fig1-7 原映射整体后错一位（image1 被误配给图注区标题）
    # 正确对应：imageN.png 内容即 Figure N（N=1..7）；image10 为旧版去卷积图，已被 codex 新图替代，弃用
    'Figure 1. Bulk': ['image1.png'],
    'Figure 2. Single-cell': ['image2.png'],
    'Figure 3. CellChat': ['image3.png'],
    'Figure 4. Spatially': ['image4.png'],
    'Figure 5. ECM-SDC4/CD44': ['image5.png'],
    'Figure 6. Nomogram': ['image6.png', 'image7.png'],
    'Figure 7. MMR-status': ['image8.png', 'image9.png'],
    'Figure 8. Immune/stromal': ['Fig8_deconvolution.png'],
    'Figure 9. Matched-null': ['codex_Fig2_robustness_nulls.png'],
    'Figure 10. The FAP-marked': ['Fig10_senescence.png'],
    'Supplementary Figure S1.': ['image11.png', 'image12.png'],
    'Supplementary Figure S2.': ['image13.png', 'image14.png'],
    'Supplementary Figure S3.': ['image15.png'],
    'Supplementary Figure S4.': ['image16.png', 'image17.png'],
    'Supplementary Figure S5.': ['image18.png'],
    'Supplementary Figure S6.': ['image19.png', 'image20.png'],
    'Supplementary Figure S7.': ['codex_S1_CMS.png'],
}
flushed = set()

def flush_figs(prefix):
    for key, imgs in FIG_MAP.items():
        if prefix.startswith(key) and key not in flushed:
            flushed.add(key)
            w = 11.0 if len(imgs) > 1 else 13.0
            for img in imgs:
                fp = os.path.join(FIGDIR, img)
                if not os.path.exists(fp):
                    print('MISSING:', fp)
                    continue
                from PIL import Image
                im = Image.open(fp)
                h = w * im.size[1] / im.size[0]
                if h > 18:
                    h = 18
                    w = 18 * im.size[0] / im.size[1]
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run()
                r.add_picture(fp, width=Cm(w), height=Cm(h))
            return

# 表格解析
table_blocks = []
cur = []
for line in tables_part.split('\n'):
    if line.startswith('--- Table '):
        if cur:
            table_blocks.append(cur)
        cur = [line]
    else:
        cur.append(line)
if cur:
    table_blocks.append(cur)
parsed_tables = []
for tb in table_blocks:
    rows = [line.split(' | ') for line in tb[1:] if line.strip()]
    parsed_tables.append({'header': tb[0], 'rows': rows})

in_refs = False
for idx, text in body_part:
    strip = text.strip()
    if idx == '117':
        in_refs = True
        add_para(strip, bold=True, space_after=8)
        continue
    if idx == '154':
        in_refs = False
    try:
        n = int(idx)
    except ValueError:
        n = 999
    if in_refs and re.match(r'^\d+\.\s', strip):
        add_para(strip, size=11, space_after=3)
        continue
    if n <= 104 or (isinstance(idx, str) and idx.startswith('3.1')):
        is_heading = any(strip.startswith(h) for h in HEADING1) or \
                     re.match(r'^\d+\.\d+\s', strip) or re.match(r'^\d+\.\s', strip)
        if is_heading:
            add_para(strip, bold=True, space_after=8)
        else:
            add_para(strip)
    elif 105 <= n <= 116:
        if strip.endswith(':'):
            add_para(strip, bold=True, space_after=4)
        else:
            add_para(strip)
    elif n == 117:
        add_para(strip, bold=True, space_after=8)
    elif 118 <= n <= 153:
        add_para(strip, size=11, space_after=3)
    elif n == 154:
        add_para(strip, bold=True, space_after=8)
    elif 155 <= n <= 168:
        if strip.startswith('Table ') or strip.startswith('Note:'):
            add_para(strip, bold=strip.startswith('Table '), size=11, space_after=4)
        else:
            add_para(strip, size=11)
    elif n >= 169:
        flush_figs(strip)
        if strip.startswith(('Figure ', 'Supplementary Figure')):
            add_para(strip, size=11, space_after=10)
        else:
            add_para(strip, size=11)

for tb in parsed_tables:
    header_line = tb['header']
    m = re.match(r'--- Table (\d+) \((\d+)x(\d+)\) ---', header_line)
    if not m:
        continue
    t = doc.add_table(rows=len(tb['rows']) + 1, cols=int(m.group(3)))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    hdr = t.rows[0]
    for ci, cell in enumerate(hdr.cells):
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run('Field')
        set_font(r, size=9, bold=True)
    for ri, row in enumerate(tb['rows']):
        for ci, cellval in enumerate(row):
            if ci >= int(m.group(3)):
                break
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(cellval.strip())
            set_font(r, size=9)

doc.save(OUT_DOCX)
print('docx saved:', OUT_DOCX)

# ---------- md 生成 ----------
def is_main_heading(t):
    return t in HEADING1 or t.startswith('Keywords:') or re.match(r'^\d+\.\s', t)

md_lines = []
in_refs = False
for idx, text in body_part:
    strip = text.strip()
    if not strip:
        continue
    if idx == '0':
        md_lines.append('# ' + strip)
    elif idx in ('1', '2', '3', '4', '5', '6'):
        md_lines.append(strip)
    elif idx == '117':
        md_lines.append('## References')
        in_refs = True
        continue
    elif in_refs and re.match(r'^\d+\.\s', strip):
        md_lines.append(strip)
        continue
    elif re.match(r'^\d+\.\d+\s', strip):
        md_lines.append('### ' + strip)
    elif is_main_heading(strip):
        md_lines.append('## ' + strip)
    elif idx.startswith('Fig') or strip.startswith(('Figure ', 'Supplementary Figure', 'Table ', 'Note:')):
        md_lines.append(strip)
    else:
        md_lines.append(strip)
    if idx == '13' or idx == '20':
        md_lines.append('')

md_out = '\n'.join(md_lines)
with open(OUT_MD, 'w', encoding='utf-8-sig') as f:
    f.write(md_out)
print('md saved:', OUT_MD, '|', len(md_lines), 'lines')
