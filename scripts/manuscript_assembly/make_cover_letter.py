# -*- coding: utf-8 -*-
"""生成 AJCR Cover Letter docx"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(2.54)
sec.left_margin = sec.right_margin = Cm(2.54)

def para(text, bold=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    r.font.size = Pt(size)
    r.font.bold = bold
    return p

para('Cover Letter', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

para('Dear Editor-in-Chief,', space_after=8)
para('We are pleased to submit our manuscript entitled "Cross-platform characterization of a FAP-associated stromal matrix state and its senescence-related transcriptional covariation in colorectal cancer" for consideration as an original research article in the American Journal of Cancer Research.', space_after=8)

para('Fibroblast activation protein (FAP) marks activated cancer-associated fibroblasts in colorectal cancer (CRC), yet bulk-tissue correlations between stromal markers and extracellular matrix (ECM) scores are vulnerable to score-gene overlap, tissue purity and platform artefacts. Using three non-overlapping gene scores (FAP13, matrix4, receptor2), we show that a FAP-marked stromal program is reproducibly associated with a collagen- and fibronectin-rich matrix state across two independent bulk cohorts (TCGA rho = 0.932; GSE39582 rho = 0.913), single-cell, and proteomic layers, and that this association exceeds matched empirical null distributions and survives ABSOLUTE purity adjustment and CMS stratification. Epithelial SDC4/CD44 receptors did not show stable co-induction with this program. We further show that the FAP-marked matrix state is enriched for senescence-associated transcriptional features across four independently constructed senescence signatures, with MKI67-adjusted robustness.', space_after=8)

para('Strengths of the study include: (1) pre-specified, non-overlapping gene scores to avoid circularity; (2) multi-cohort, multi-layer independent validation (bulk, single-cell, spatial, proteomic); (3) three matched empirical null distributions (10,000 draws each); (4) explicit treatment of tissue purity and CMS stratification; (5) honest reporting of null results (receptor co-induction, prognostic value). We believe this rigorous computational study is well within the scope of AJCR, which welcomes bioinformatics analyses of public cancer data.', space_after=8)

para('Suggested reviewers (4):', bold=True, space_after=4)
for i, name in enumerate([
    '1. [Name, MD/PhD], [Institution, City, Country], [email] — expertise: cancer-associated fibroblasts / tumor microenvironment',
    '2. [Name, MD/PhD], [Institution, City, Country], [email] — expertise: extracellular matrix remodelling / colorectal cancer',
    '3. [Name, MD/PhD], [Institution, City, Country], [email] — expertise: cellular senescence / SASP in cancer',
    '4. [Name, MD/PhD], [Institution, City, Country], [email] — expertise: bioinformatics / TCGA multi-omics analysis',
]):
    para('   ' + name, size=10, space_after=2)

para('Non-preferred reviewers (2):', bold=True, space_after=4)
para('   1. [Name], [Institution] — reason: [e.g., direct competitor]', size=10, space_after=2)
para('   2. [Name], [Institution] — reason: [e.g., known conflict of interest]', size=10, space_after=4)

para('All authors have read and approved the manuscript, and no part of this work has been published or is under consideration elsewhere. The authors declare no competing interests. This study used only publicly available, de-identified data and required no ethics approval.', space_after=8)

para('Thank you for your consideration.', space_after=8)
para('Sincerely,', space_after=4)
para('Jinglin Wang, MD', space_after=2)
para('Department of Geriatrics, The Affiliated Hospital of Guizhou Medical University, Guiyang, China', size=10, space_after=2)
para('Corresponding author: Juan Yang, MD, PhD (yj63yj63@163.com)', space_after=8)

doc.save(r'C:/Users/Spica/WorkBuddy/Claw/deliverables/FAP_SDC4_CD44_v5.0/投稿图片/AJCR_CoverLetter_2026-08-08.docx')
print('Cover Letter 已生成')
